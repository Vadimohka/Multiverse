from __future__ import annotations

import ast
import asyncio
import base64
import csv
import hashlib
import hmac
import io
import json
import re
from collections.abc import Mapping
from datetime import UTC, date, datetime, timedelta
from datetime import time as datetime_time
from fnmatch import fnmatchcase
from html import escape
from pathlib import Path
from typing import Any
from urllib.parse import parse_qsl, urlencode, urljoin, urlsplit, urlunsplit
from zoneinfo import ZoneInfo

import httpx
from bs4 import BeautifulSoup
from lxml import etree
from lxml import html as lxml_html
from openpyxl import Workbook, load_workbook

from .egress import BrowserEgressGuard, EgressPolicy, default_resolver, request_with_egress_policy
from .normalizers import normalize_currency, normalize_number, normalize_term, parse_rate_expression
from .redaction import redact_artifact_bytes, redact_text
from .transport import FetchPolicy, fetch_attempts, request_with_policy
from .types import ExecutionContext


class ManualTriggerNode:
    type = "manual_trigger"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        return {"data": inputs or context.variables, **(inputs or context.variables)}


class HTTPRequestNode:
    type = "http_request"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        url = render_template(str(config.get("url") or "{{source.url}}"), context, inputs)
        if not url:
            raise ValueError("HTTP Request: URL не задан")
        source_settings = context.variables.get("source", {}).get("settings", {})
        profile = source_settings.get("profile", {}) if isinstance(source_settings, dict) else {}
        # Source Profiler runs in the UI before a source is created. Reuse its
        # verdict so public pages rendered by JavaScript are not parsed as an
        # empty static shell.
        needs_browser = bool(profile.get("requires_javascript")) or source_settings.get("fetch_mode") == "PLAYWRIGHT"
        if needs_browser and not config.get("_force_http"):
            return await BrowserOpenNode().execute(
                context, inputs, {**config, "url": url, "http_fallback": True, "_force_http": True}
            )
        method = str(config.get("method", "GET")).upper()
        policy = FetchPolicy.from_config(config)
        headers = render_object(config.get("headers") or {}, context, inputs)
        cookies = render_object(config.get("cookies") or {}, context, inputs)
        # ``httpx`` replaces the query string embedded in ``url`` when
        # ``params={}`` is passed.  An empty editor value must therefore mean
        # "preserve the selected Source URL", not "strip its query".  This is
        # especially important for public JSON endpoints whose required query
        # parameters live in the Source URL.
        query_params = render_object(config.get("query_params") or {}, context, inputs)
        if not query_params:
            query_params = None
        json_body = render_object(config.get("json_body") or {}, context, inputs)
        resolver = config.get("egress_resolver")
        async with httpx.AsyncClient(follow_redirects=False, cookies=cookies) as client:
            response = await request_with_egress_policy(
                client,
                method,
                url,
                policy,
                egress_policy=EgressPolicy.from_config(config),
                resolver=resolver or default_resolver,
                request_fn=request_with_policy,
                headers=headers,
                params=query_params,
                json=json_body or None,
            )
            response.raise_for_status()
        payload = await response_payload(context, response)
        payload["fetch_attempts"] = fetch_attempts(response)
        payload["redirect_chain"] = response.extensions.get("redirect_chain", [])
        return payload


class BrowserOpenNode:
    type = "browser_open"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        url = render_template(str(config.get("url") or "{{source.url}}"), context, inputs)
        if not url:
            raise ValueError("Browser Open: URL не задан")
        egress_policy = EgressPolicy.from_config(config)
        resolver = config.get("egress_resolver") or default_resolver
        egress_policy.validate_url(url, resolver=resolver)
        timeout_ms = int(float(config.get("timeout", 45)) * 1000)
        network: list[dict[str, Any]] = []
        try:
            from playwright.async_api import async_playwright

            async with async_playwright() as playwright:
                profile = context.capabilities.get("browser_profile", {})
                proxy = profile.get("proxy") if isinstance(profile, dict) else None
                browser = await playwright.chromium.launch(headless=True, proxy=proxy or None)
                browser_context = await browser.new_context(
                    **browser_context_options(profile if isinstance(profile, dict) else {}, config)
                )
                egress_guard = BrowserEgressGuard(egress_policy, resolver=resolver)
                await egress_guard.install(browser_context)
                configured_cookies = config.get("browser_cookies") or profile.get("cookies") or []
                if isinstance(configured_cookies, dict):
                    origin = urlunsplit((*urlsplit(url)[:2], "", "", ""))
                    configured_cookies = [
                        {"name": name, "value": str(value), "url": origin}
                        for name, value in configured_cookies.items()
                    ]
                if configured_cookies:
                    await browser_context.add_cookies(configured_cookies)
                page = await browser_context.new_page()
                if config.get("capture_network", True):
                    async def on_response(response: Any) -> None:
                        content_type = response.headers.get("content-type", "")
                        if "json" in content_type or response.request.resource_type in {"xhr", "fetch"}:
                            try:
                                body = await response.json()
                            except Exception:
                                return
                            network.append({"url": response.url, "status": response.status, "content_type": content_type, "body": body})
                    page.on("response", on_response)
                await page.goto(url, wait_until=str(config.get("wait_until", "networkidle")), timeout=timeout_ms)
                egress_guard.assert_safe()
                for action in config.get("actions", []):
                    await perform_browser_action(page, action, timeout_ms)
                    egress_guard.assert_safe()
                tab_descriptors = await discover_tab_descriptors(page) if config.get("tabs_enabled", False) else []
                rendered_html = await collect_paginated_html(page, config, timeout_ms, start_url=page.url, egress_guard=egress_guard)
                screenshot = await page.screenshot(full_page=bool(config.get("full_page", True)), type="png")
                title = await page.title()
                final_url = page.url
                final_target = egress_policy.validate_url(final_url, resolver=resolver)
                egress_guard.assert_safe()
                await browser_context.close()
                await browser.close()
            raw = rendered_html.encode("utf-8")
            html_artifact = await store_artifact(context, raw, "text/html", final_url, "rendered.html", "rendered_html")
            screenshot_artifact = await store_artifact(context, screenshot, "image/png", final_url, "screenshot.png", "screenshot")
            artifacts = [html_artifact, screenshot_artifact]
            if network:
                network_artifact = await store_artifact(
                    context,
                    json.dumps(network, ensure_ascii=False).encode("utf-8"),
                    "application/json",
                    final_url,
                    "network.json",
                    "network_capture",
                )
                artifacts.append(network_artifact)
            return {
                "url": final_url,
                "title": title,
                "body": rendered_html,
                "html": rendered_html,
                "text": BeautifulSoup(rendered_html, "lxml").get_text("\n", strip=True),
                "network": network,
                "tab_count": len(tab_descriptors),
                "tab_labels": [item.get("text") for item in tab_descriptors],
                "artifacts": artifacts,
                "browser_mode": "PLAYWRIGHT",
                "redirect_chain": egress_guard.redirect_chain or [{
                    "hop": 0,
                    "requested_url": url,
                    "location": final_url,
                    "resolved_addresses": final_target["addresses"],
                }],
            }
        except ImportError as exc:
            context.log("WARNING", "Playwright не установлен; использован HTTP fallback", error=str(exc))
        except Exception as exc:
            if not config.get("http_fallback", True):
                raise
            context.log("WARNING", "Browser Open завершился ошибкой; использован HTTP fallback", error=str(exc))
        fallback = await HTTPRequestNode().execute(context, inputs, {**config, "url": url, "_force_http": True})
        fallback["html"] = fallback.get("body")
        fallback["browser_mode"] = "HTTP_FALLBACK"
        return fallback


async def perform_browser_action(page: Any, action: dict[str, Any], timeout_ms: int) -> None:
    kind = str(action.get("type") or "")
    selector = action.get("selector")
    known = {"click", "fill", "select", "hover", "press", "wait", "wait_for", "scroll", "javascript"}
    if kind not in known:
        raise ValueError(f"Unknown browser action: {kind or '<empty>'}")
    if kind in {"click", "fill", "select", "hover", "wait_for"} and not selector:
        raise ValueError(f"Browser action {kind} requires selector")
    if kind == "javascript" and not str(action.get("script") or "").strip():
        raise ValueError("Browser action javascript requires script")
    if kind == "click":
        await page.locator(selector).first.click(timeout=timeout_ms)
    elif kind == "fill":
        await page.locator(selector).first.fill(str(action.get("value", "")), timeout=timeout_ms)
    elif kind == "select":
        await page.locator(selector).first.select_option(str(action.get("value", "")), timeout=timeout_ms)
    elif kind == "hover":
        await page.locator(selector).first.hover(timeout=timeout_ms)
    elif kind == "press":
        await page.locator(selector or "body").press(str(action.get("value", "Enter")), timeout=timeout_ms)
    elif kind == "wait":
        await page.wait_for_timeout(int(float(action.get("seconds", 1)) * 1000))
    elif kind == "wait_for":
        await page.locator(selector).first.wait_for(timeout=timeout_ms)
    elif kind == "scroll":
        await page.evaluate("window.scrollBy(0, arguments[0])", int(action.get("pixels", 1000)))
    elif kind == "javascript":
        await page.evaluate(str(action.get("script", "")))


def browser_context_options(profile: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    options: dict[str, Any] = {
        "viewport": config.get("viewport") or profile.get("viewport") or {"width": 1440, "height": 900},
        "locale": str(config.get("locale") or profile.get("locale") or "ru-RU"),
        "timezone_id": str(config.get("timezone") or profile.get("timezone") or "Europe/Minsk"),
    }
    user_agent = config.get("user_agent") or profile.get("user_agent")
    storage_state = config.get("storage_state") or profile.get("storage_state")
    if user_agent:
        options["user_agent"] = str(user_agent)
    if storage_state:
        options["storage_state"] = storage_state
    return options


async def collect_paginated_html(
    page: Any,
    config: dict[str, Any],
    timeout_ms: int,
    start_url: str | None = None,
    egress_guard: BrowserEgressGuard | None = None,
) -> str:
    """Collect every semantic tab and every page exposed by its paginator.

    The controls are discovered from ARIA/data-toggle semantics rather than
    site-specific classes or URLs.  Each tab is revisited from the original
    listing URL so a tab that changes the URL cannot hide subsequent tabs.
    """
    tabs_enabled = bool(config.get("tabs_enabled", False))
    initial_url = start_url or page.url
    descriptors = await discover_tab_descriptors(page) if tabs_enabled else []
    if descriptors:
        sections: list[str] = []
        wait_ms = int(config.get("tabs_wait_ms") or config.get("pagination_wait_ms") or 500)

        async def visit(path: list[dict[str, str]]) -> None:
            try:
                await page.goto(initial_url, wait_until="domcontentloaded", timeout=timeout_ms)
                if egress_guard:
                    egress_guard.assert_safe()
                for control in path:
                    await click_tab_descriptor(page, control, timeout_ms)
                    await page.wait_for_timeout(wait_ms)
                    if egress_guard:
                        egress_guard.assert_safe()
                visible = await discover_tab_descriptors(page)
                current_depth = int(path[-1].get("scope_depth") or 0)
                path_keys = {tab_descriptor_key(item) for item in path}
                children = [item for item in visible if int(item.get("scope_depth") or 0) > current_depth and tab_descriptor_key(item) not in path_keys]
                if children and len(path) < int(config.get("tabs_max_depth") or 4):
                    for child in children:
                        await visit(path + [child])
                else:
                    sections.append(await collect_current_paginated_html(page, config, timeout_ms, egress_guard=egress_guard))
            except Exception:
                return

        for descriptor in descriptors:
            await visit([descriptor])
        if sections:
            return merge_rendered_sections(sections)
    return await collect_current_paginated_html(page, config, timeout_ms, egress_guard=egress_guard)


async def discover_tab_descriptors(page: Any) -> list[dict[str, str]]:
    """Return visible, semantic tab controls with stable attributes."""
    script = """
    () => Array.from(document.querySelectorAll(
      '[role="tab"], [aria-controls], [data-toggle="tab"], [data-bs-toggle="tab"]'
    )).filter((el) => {
      const rect = el.getBoundingClientRect();
      const style = getComputedStyle(el);
      return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
    }).map((el) => ({
      id: el.id || '', role: el.getAttribute('role') || '',
      aria_controls: el.getAttribute('aria-controls') || '',
      href: el.getAttribute('href') || '',
      text: (el.innerText || el.textContent || '').trim().replace(/\\s+/g, ' ').slice(0, 160),
      scope_depth: (() => { let depth = 0, node = el.parentElement; while (node) { if (node.getAttribute && node.getAttribute('role') === 'tabpanel') depth += 1; node = node.parentElement; } return depth; })()
    }))
    """
    values = await page.evaluate(script)
    result: list[dict[str, str]] = []
    seen: set[tuple[str, str, str, str, str]] = set()
    for item in values or []:
        if not isinstance(item, dict) or not item.get("text"):
            continue
        key = (str(item.get("id") or ""), str(item.get("aria_controls") or ""), str(item.get("href") or ""), str(item.get("text") or ""), str(item.get("scope_depth") or "0"))
        if key in seen:
            continue
        seen.add(key)
        result.append({key_name: str(item.get(key_name) or "") for key_name in ("id", "role", "aria_controls", "href", "text", "scope_depth")})
    return result


def tab_descriptor_key(item: dict[str, str]) -> tuple[str, str, str, str, str]:
    return (item.get("id", ""), item.get("aria_controls", ""), item.get("href", ""), item.get("text", ""), item.get("scope_depth", "0"))


async def click_tab_descriptor(page: Any, descriptor: dict[str, str], timeout_ms: int) -> None:
    selectors: list[str] = []
    if descriptor.get("id"):
        selectors.append(f"#{css_escape(descriptor['id'])}")
    if descriptor.get("aria_controls"):
        selectors.append(f"[aria-controls='{css_escape(descriptor['aria_controls'])}']")
    if descriptor.get("href"):
        selectors.append(f"[href='{css_escape(descriptor['href'])}']")
    for selector in selectors:
        locator = page.locator(selector).first
        if await locator.count() and await locator.is_visible():
            await locator.click(timeout=timeout_ms)
            return
    locator = page.get_by_text(descriptor.get("text", ""), exact=True).first
    if await locator.count() and await locator.is_visible():
        await locator.click(timeout=timeout_ms)
        return
    raise ValueError("Semantic tab control is no longer available")


def css_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("'", "\\'")


async def collect_current_paginated_html(
    page: Any,
    config: dict[str, Any],
    timeout_ms: int,
    egress_guard: BrowserEgressGuard | None = None,
) -> str:
    """Collect the current tab through visible pagination controls."""
    if not config.get("pagination_enabled"):
        return await page.content()
    max_pages = min(max(int(config.get("pagination_max_pages") or 25), 1), 500)
    selector = str(config.get("pagination_next_selector") or "li[aria-label='Next page'] a")
    rendered_pages: list[str] = []
    seen_signatures: set[str] = set()
    for index in range(max_pages):
        document = await page.content()
        rendered_pages.append(document)
        signature = await page_text_signature(page)
        if signature in seen_signatures:
            break
        seen_signatures.add(signature)
        # ``max_pages`` is a fetch budget, not a number of navigation
        # attempts.  After retaining the last requested page, do not advance
        # the live tab one more time; doing so used to leave provenance on an
        # uncollected third page for a two-page traversal.
        if index + 1 >= max_pages:
            break
        candidates = page.locator(selector)
        next_link = None
        for index in range(await candidates.count()):
            candidate = candidates.nth(index)
            if await candidate.is_visible():
                next_link = candidate
                break
        if next_link is None:
            # Some client-side pagers expose an empty-href anchor without an
            # aria label.  Use only visible conventional next-page captions as
            # a semantic fallback; no site URL or CSS class is required.
            candidates = page.locator("a, button")
            for index in range(await candidates.count()):
                candidate = candidates.nth(index)
                if (await candidate.is_visible()) and (await candidate.inner_text()).strip() in {"»", ">", "Next", "Следующая"}:
                    next_link = candidate
                    break
        if next_link is None or not await next_link.is_enabled():
            break
        try:
            await next_link.click(timeout=timeout_ms)
            await page.wait_for_timeout(int(config.get("pagination_wait_ms") or 500))
            await page.wait_for_load_state("networkidle", timeout=min(timeout_ms, 15000))
            if egress_guard:
                egress_guard.assert_safe()
        except Exception:
            break
        if await page_text_signature(page) == signature:
            break
    if len(rendered_pages) == 1:
        return rendered_pages[0]
    return merge_rendered_sections(rendered_pages)


def merge_rendered_sections(documents: list[str]) -> str:
    mains: list[str] = []
    bodies: list[str] = []
    for document in documents:
        soup = BeautifulSoup(document, "lxml")
        main = soup.select_one("main")
        body = soup.body
        if main:
            mains.append(str(main))
        elif body:
            bodies.append(str(body))
        else:
            bodies.append(str(soup))
    if mains:
        return "<html><body><main>" + "\n".join(mains) + "</main></body></html>"
    return "<html><body>" + "\n".join(bodies) + "</body></html>"


async def page_text_signature(page: Any) -> str:
    locator = page.locator("main")
    if await locator.count():
        text = await locator.first.inner_text()
    else:
        text = await page.locator("body").inner_text()
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


class DownloadFileNode:
    type = "download_file"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        url = render_template(str(config.get("url") or "{{source.url}}"), context, inputs)
        source_settings = context.variables.get("source", {}).get("settings", {})
        if source_settings.get("document_storage_key"):
            if context.artifact_storage is None:
                raise ValueError("Document source requires artifact storage")
            content_type = str(source_settings.get("document_content_type") or "application/octet-stream")
            filename = str(source_settings.get("document_filename") or "document")
            content = await context.artifact_storage.get_bytes(
                str(source_settings.get("document_bucket") or "raw"),
                str(source_settings["document_storage_key"]),
                str(source_settings.get("document_storage_backend") or "S3"),
            )
            artifact = await store_artifact(context, content, content_type, url, filename, "raw_document")
            return {"url": url, "filename": filename, "content_type": content_type, "content_base64": base64.b64encode(content).decode("ascii"), "size": len(content), "sha256": artifact["sha256"], "artifact": artifact}
        policy = FetchPolicy.from_config(config)
        cookies = render_object(config.get("cookies") or {}, context, inputs)
        resolver = config.get("egress_resolver")
        async with httpx.AsyncClient(follow_redirects=False, cookies=cookies) as client:
            response = await request_with_egress_policy(
                client,
                "GET",
                url,
                policy,
                egress_policy=EgressPolicy.from_config(config),
                resolver=resolver or default_resolver,
                request_fn=request_with_policy,
                headers=render_object(config.get("headers") or {}, context, inputs),
            )
            response.raise_for_status()
        content_type = response.headers.get("content-type", "application/octet-stream")
        filename = filename_from_response(response)
        artifact = await store_artifact(context, response.content, content_type, str(response.url), filename, "raw_document")
        return {
            "url": str(response.url), "filename": filename, "content_type": content_type,
            "content_base64": base64.b64encode(response.content).decode("ascii"), "size": len(response.content),
            "sha256": artifact["sha256"], "artifact": artifact,
            "fetch_attempts": fetch_attempts(response),
            "redirect_chain": response.extensions.get("redirect_chain", []),
        }


class FollowLinksNode:
    type = "follow_links"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        collection = find_value(inputs, str(config.get("input_collection", "records")))
        base_url = str(find_value(inputs, "url") or context.variables.get("source", {}).get("url", ""))
        # Backwards-compatible listing mode: turn selected HTML links into the
        # same explicit parent collection before fan-out.
        if not isinstance(collection, list):
            html = find_value(inputs, str(config.get("input_path", "html"))) or find_value(inputs, "body")
            soup = BeautifulSoup(str(html or ""), "lxml")
            collection = [{str(config.get("url_field", "url")): urljoin(base_url, element.get("href", ""))}
                          for element in soup.select(str(config.get("selector", "a[href]"))) if element.get("href")]
        url_field = str(config.get("url_field", "url"))
        limit = max(0, int(config.get("max_pages") or len(collection) or 20))
        frontier = build_url_frontier(
            collection,
            base_url=base_url,
            origin_url=base_url,
            url_path=url_field,
            config=config,
            limit=limit or len(collection),
        )
        parents = []
        for candidate in frontier:
            parent = dict(candidate["item"])
            parent[url_field] = candidate["url"]
            parents.append(parent)

        concurrency = min(max(int(config.get("concurrency") or 3), 1), 20)
        fetch_policy = FetchPolicy.from_config(config)
        merge_mode = str(config.get("merge_mode", "MERGE_PARENT_CHILD"))
        policy = str(config.get("error_policy", "CONTINUE"))
        progress: list[dict[str, Any]] = []
        records: list[dict[str, Any]] = []
        failures: list[dict[str, Any]] = []
        lock = asyncio.Lock()
        semaphore = asyncio.Semaphore(concurrency)

        headers = render_object(config.get("headers") or {}, context, inputs)
        cookies = render_object(config.get("cookies") or {}, context, inputs)
        egress_policy = EgressPolicy.from_config(config)
        resolver = config.get("egress_resolver") or default_resolver
        async with httpx.AsyncClient(follow_redirects=False, headers=headers, cookies=cookies) as client:
            async def fetch(parent: dict[str, Any]) -> None:
                url = str(parent[url_field])
                response: httpx.Response | None = None
                message = ""
                exception_attempts: list[dict[str, Any]] = []
                exception_redirect_chain: list[dict[str, Any]] = []
                async with semaphore:
                    try:
                        response = await request_with_egress_policy(
                            client,
                            "GET",
                            url,
                            fetch_policy,
                            egress_policy=egress_policy,
                            resolver=resolver,
                            request_fn=request_with_policy,
                        )
                        response.raise_for_status()
                    except Exception as exc:
                        message = str(exc)
                        exception_attempts = getattr(exc, "fetch_attempts", [])
                        exception_redirect_chain = getattr(exc, "redirect_chain", [])
                    if response is None or not response.is_success:
                        async with lock:
                            attempts = fetch_attempts(response) if response is not None else exception_attempts
                            redirect_chain = response.extensions.get("redirect_chain", []) if response is not None else exception_redirect_chain
                            failures.append({"url": url, "error": message, "fetch_attempts": attempts, "redirect_chain": redirect_chain})
                            progress.append({"url": url, "status": "FAILED", "error": message, "fetch_attempts": attempts, "redirect_chain": redirect_chain})
                        return
                    child: dict[str, Any] = {"url": str(response.url), "status_code": response.status_code, "body": response.text}
                    soup = BeautifulSoup(response.text, "lxml")
                    for mapping in config.get("detail_fields", []):
                        name = mapping.get("target") or mapping.get("name")
                        if not name:
                            continue
                        selector = mapping.get("selector") or mapping.get("source_path")
                        element = soup.select_one(str(selector)) if selector else None
                        child[name] = element.get(mapping.get("attribute")) if element and mapping.get("attribute") else element.get_text(" ", strip=True) if element else mapping.get("default")
                    table_config = config.get("detail_table")
                    table_rows = []
                    if isinstance(table_config, dict) and table_config.get("selector"):
                        parsed_table = await ParseTableNode().execute(context, {"html": response.text}, table_config)
                        table_rows = parsed_table.get("records") or []
                    rows = table_rows or [{}]
                    async with lock:
                        for table_row in rows:
                            detail = {**child, **table_row}
                            row = parent if merge_mode == "PARENT_ONLY" else detail if merge_mode == "CHILD_ONLY" else {**parent, **detail}
                            records.append(row)
                        progress.append({"url": url, "status": "SUCCESS", "status_code": response.status_code, "detail_rows": len(table_rows), "fetch_attempts": fetch_attempts(response), "redirect_chain": response.extensions.get("redirect_chain", [])})

            await asyncio.gather(*(fetch(parent) for parent in parents))
        if failures and policy == "FAIL_FAST":
            raise ValueError(f"Follow Links failed for {failures[0]['url']}: {failures[0]['error']}")
        return {"records": records, "links": [parent[url_field] for parent in parents], "count": len(records),
                "progress": progress, "errors": failures, "partial": bool(failures and records)}


class PaginationNode:
    type = "pagination"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        template = str(config.get("url_template") or "")
        if not template:
            raise ValueError("Pagination: url_template не задан")
        current = int(config.get("start") or 1)
        step = int(config.get("step") or 1)
        pages: list[dict[str, Any]] = []
        policy = FetchPolicy.from_config(config)
        egress_policy = EgressPolicy.from_config(config)
        resolver = config.get("egress_resolver") or default_resolver
        async with httpx.AsyncClient(follow_redirects=False, timeout=policy.timeout) as client:
            for _ in range(min(int(config.get("max_pages") or 10), 1000)):
                url = render_template(template, context, {**inputs, "page": current, "offset": current})
                response = await request_with_egress_policy(
                    client,
                    "GET",
                    url,
                    policy,
                    egress_policy=egress_policy,
                    resolver=resolver,
                    request_fn=request_with_policy,
                )
                if not response.is_success:
                    break
                body = response.text
                stop_selector = str(config.get("stop_selector") or "")
                if stop_selector and not BeautifulSoup(body, "lxml").select(stop_selector):
                    break
                pages.append({"url": str(response.url), "status_code": response.status_code, "body": body, "fetch_attempts": fetch_attempts(response), "redirect_chain": response.extensions.get("redirect_chain", [])})
                current += step
        return {"pages": pages, "count": len(pages)}


class CrawlLinksNode:
    """Fetch a collection of links and turn every detail page into a record.

    The workflow engine deliberately has no hidden subgraph semantics.  This node is
    therefore the explicit, observable fan-out/fan-in primitive for collection
    crawls: list item -> HTTP request -> configured detail extraction -> record.
    """

    type = "crawl_links"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        config = self._effective_config(context, inputs, config)
        headers = {"Accept-Language": "ru-RU,ru;q=0.9,en;q=0.5", "User-Agent": "Mozilla/5.0 (compatible; ParserStudio/1.0)"}
        headers.update(render_object(config.get("headers") or {}, context, inputs))
        cookies = render_object(config.get("cookies") or {}, context, inputs)
        client = httpx.AsyncClient(follow_redirects=False, headers=headers, cookies=cookies)
        resume_token = str(inputs.get("resume_token") or config.get("resume_token") or "").strip()
        try:
            if resume_token:
                resume_urls = decode_crawl_resume_token(resume_token, context, config)
                listing = [{"url": url} for url in resume_urls]
                listing_url = resume_urls[0] if resume_urls else str(config.get("listing_url") or "")
                context.variables["_crawl_listing_diagnostics"] = {
                    "fetch_mode": "RESUME",
                    "resumed_urls": len(resume_urls),
                }
            else:
                listing, listing_url = await self._load_listing(context, inputs, config, client=client)
        except Exception:
            await close_http_client(client)
            raise
        listing_diagnostics = context.variables.pop("_crawl_listing_diagnostics", {})
        items = self._listing_items(listing, config)
        configured_base = str(config.get("base_url") or find_value(inputs, "url") or context.variables.get("source", {}).get("base_url", "") or context.variables.get("source", {}).get("url", ""))
        listing_parts = urlsplit(listing_url)
        listing_origin = urlunsplit((listing_parts.scheme, listing_parts.netloc, "/", "", "")) if listing_parts.scheme and listing_parts.netloc else ""
        base_url = configured_base or listing_origin or listing_url
        pattern_text = str(config.get("url_pattern") or "").strip()
        pattern = re.compile(pattern_text, re.I) if pattern_text else None
        url_path = str(config.get("url_path") or "url")
        maximum = min(
            max(int(config.get("max_items") or 5000), 1),
            max(int(config.get("max_pages") or config.get("max_items") or 5000), 1),
            5000,
        )
        candidates = build_url_frontier(
            items,
            base_url=base_url,
            origin_url=listing_url,
            url_path=url_path,
            config=config,
            limit=maximum,
        )
        for candidate in candidates:
            match = pattern.search(candidate["url"]) if pattern else None
            if match:
                candidate.update({key: value for key, value in match.groupdict().items() if value is not None})
            candidate["record_id"] = (
                candidate.get("record_id")
                or (candidate.get("item") or {}).get("record_id")
                or (match.group(1) if match and match.groups() else match.group(0) if match else None)
                or hashlib.sha256(candidate["url"].encode("utf-8")).hexdigest()[:20]
            )
            candidate["depth"] = 1

        concurrency = min(max(int(config.get("concurrency") or 3), 1), 20)
        delay_ms = max(int(config.get("delay_ms") or 400), 0)
        fetch_policy = FetchPolicy.from_config(config)
        egress_policy = EgressPolicy.from_config(config)
        resolver = config.get("egress_resolver") or default_resolver
        semaphore = asyncio.Semaphore(concurrency)
        records: list[dict[str, Any]] = []
        errors: list[dict[str, Any]] = []
        detail_diagnostics: list[dict[str, Any]] = []
        record_lock = asyncio.Lock()
        visited_urls = {candidate["url"] for candidate in candidates}
        all_candidates = list(candidates)
        max_depth = min(max(int(config.get("max_depth") or 1), 1), 20)
        recursive_selector = str(config.get("recursive_link_selector") or "").strip()

        async def crawl(candidate: dict[str, Any]) -> list[dict[str, Any]]:
                if context.cancelled:
                    return []
                async with semaphore:
                    detail_url = candidate["url"]
                    detail_html = ""
                    artifact_content = b""
                    artifact_content_type = "text/html"
                    error: Exception | None = None
                    response: httpx.Response | None = None
                    detail_request = config.get("detail_request") if isinstance(config.get("detail_request"), dict) else None
                    if detail_request:
                        detail_scope = {**inputs, **candidate, "item": candidate.get("item") or {}}
                        request_url = render_template(str(detail_request.get("url") or candidate["url"]), context, detail_scope)
                        request_params = render_object(detail_request.get("query_params") or {}, context, detail_scope)
                        try:
                            response = await request_with_egress_policy(
                                client,
                                str(detail_request.get("method") or "GET").upper(),
                                request_url,
                                fetch_policy,
                                egress_policy=egress_policy,
                                resolver=resolver,
                                request_fn=request_with_policy,
                                params=request_params,
                            )
                            response.raise_for_status()
                        except Exception as exc:
                            error = exc
                        if response is None or not response.is_success:
                            async with record_lock:
                                errors.append({
                                    "url": candidate["url"],
                                    "record_id": candidate["record_id"],
                                    "error": str(error or "Detail API request failed"),
                                    "fetch_attempts": fetch_attempts(response) if response is not None else getattr(error, "fetch_attempts", []),
                                    "redirect_chain": response.extensions.get("redirect_chain", []) if response is not None else getattr(error, "redirect_chain", []),
                                })
                            return []
                        try:
                            detail_payload = response.json()
                        except Exception as exc:
                            async with record_lock:
                                errors.append({"url": candidate["url"], "record_id": candidate["record_id"], "error": f"Detail API returned invalid JSON: {exc}"})
                            return []
                        not_found_path = str(detail_request.get("not_found_path") or "")
                        if not_found_path and bool(find_value(detail_payload, not_found_path)):
                            async with record_lock:
                                errors.append({
                                    "url": candidate["url"],
                                    "record_id": candidate["record_id"],
                                    "error": "Detail API reported that the record was not found",
                                })
                            return []
                        candidate["detail_response"] = detail_payload
                        candidate["_detail_request_url"] = str(response.url)
                        detail_url = candidate["url"]
                        detail_html = str(find_value(detail_payload, str(detail_request.get("html_path") or "")) or "")
                        artifact_content = response.content
                        artifact_content_type = response.headers.get("content-type", "application/json")
                    elif self._detail_uses_browser(context, config):
                        rendered: dict[str, Any] | None = None
                        for attempt in range(fetch_policy.retries + 1):
                            try:
                                # Explicit browser transport must not depend on
                                # a successful direct HTTP probe. Some sites
                                # intentionally reject non-browser requests.
                                rendered = await BrowserOpenNode().execute(context, inputs, {
                                    "url": candidate["url"], "wait_until": config.get("detail_wait_until", "networkidle"),
                                    "timeout": config.get("request_timeout", 45), "headers": headers,
                                    "capture_network": False, "full_page": False, "http_fallback": False,
                                })
                                break
                            except Exception as exc:
                                error = exc
                                if attempt < fetch_policy.retries:
                                    await asyncio.sleep(min(0.5 * (attempt + 1), 2))
                        if rendered is None:
                            async with record_lock:
                                errors.append({"url": candidate["url"], "record_id": candidate["record_id"], "error": str(error or "Browser request failed")})
                            return []
                        detail_url = str(rendered.get("url") or candidate["url"])
                        detail_html = str(rendered.get("html") or rendered.get("body") or "")
                        artifact_content = detail_html.encode("utf-8")
                    else:
                        try:
                            response = await request_with_egress_policy(
                                client,
                                "GET",
                                candidate["url"],
                                fetch_policy,
                                egress_policy=egress_policy,
                                resolver=resolver,
                                request_fn=request_with_policy,
                            )
                            response.raise_for_status()
                        except Exception as exc:  # request failures are reported per item, not hidden
                            error = exc
                        if response is None or not response.is_success:
                            async with record_lock:
                                errors.append({
                                    "url": candidate["url"],
                                    "record_id": candidate["record_id"],
                                    "error": str(error or "HTTP request failed"),
                                    "fetch_attempts": fetch_attempts(response) if response is not None else getattr(error, "fetch_attempts", []),
                                    "redirect_chain": response.extensions.get("redirect_chain", []) if response is not None else getattr(error, "redirect_chain", []),
                                })
                            return []
                        detail_url = str(response.url)
                        detail_html = response.text
                        artifact_content = response.content
                        artifact_content_type = response.headers.get("content-type", "text/html")
                    artifact: dict[str, Any] | None = None
                    if config.get("save_artifacts", True):
                        artifact_id = re.sub(r"[^a-zA-Z0-9._-]+", "_", str(candidate["record_id"]))
                        artifact = await store_artifact(context, artifact_content, artifact_content_type, detail_url, f"{artifact_id}.json" if "json" in artifact_content_type else f"{artifact_id}.html", "raw_article")
                    fetched_candidate = {**candidate, "fetched_at": datetime.now(UTC).isoformat()}
                    record = extract_article_record(detail_html, detail_url, fetched_candidate, config, artifact)
                    if response is not None:
                        async with record_lock:
                            detail_diagnostics.append({
                                "url": detail_url,
                                "request_url": candidate.get("_detail_request_url") or detail_url,
                                "status_code": response.status_code,
                                "fetch_attempts": fetch_attempts(response),
                                "redirect_chain": response.extensions.get("redirect_chain", []),
                            })
                    async with record_lock:
                        records.append(record)
                    if delay_ms:
                        await asyncio.sleep(delay_ms / 1000)
                    if not recursive_selector or int(candidate["depth"]) >= max_depth:
                        return []
                    detail_items = [
                        {"url": element.get("href")}
                        for element in BeautifulSoup(detail_html, "lxml").select(recursive_selector)
                        if element.get("href")
                    ]
                    discovered = build_url_frontier(
                        detail_items,
                        base_url=detail_url,
                        origin_url=listing_url,
                        url_path="url",
                        config=config,
                        limit=maximum,
                    )
                    children: list[dict[str, Any]] = []
                    async with record_lock:
                        for child in discovered:
                            if child["url"] in visited_urls or len(all_candidates) + len(children) >= maximum:
                                continue
                            visited_urls.add(child["url"])
                            child["record_id"] = hashlib.sha256(child["url"].encode()).hexdigest()[:20]
                            child["depth"] = int(candidate["depth"]) + 1
                            children.append(child)
                    return children

        try:
            pending = candidates
            while pending:
                discovered_groups = await asyncio.gather(*(crawl(candidate) for candidate in pending))
                remaining = max(0, maximum - len(all_candidates))
                pending = [child for group in discovered_groups for child in group][:remaining]
                all_candidates.extend(pending)
        finally:
            await close_http_client(client)

        records.sort(key=lambda row: (str(row.get("published_at", "")), str(row.get("record_id", ""))), reverse=True)
        failed_urls = {str(item["url"]) for item in errors}
        completed_urls = [candidate["url"] for candidate in all_candidates if candidate["url"] not in failed_urls]
        next_resume_token = (
            encode_crawl_resume_token(sorted(failed_urls), context, config)
            if failed_urls
            else None
        )
        error_policy = str(config.get("error_policy") or "CONTINUE").upper()
        if errors and error_policy in {"FAIL", "FAIL_FAST"}:
            raise ValueError(f"Crawl failed for {errors[0]['url']}: {errors[0]['error']}")
        if error_policy == "REQUIRE_MINIMUM":
            minimum = max(int(config.get("minimum_successful_records") or 1), 0)
            if len(records) < minimum:
                raise ValueError(f"Crawl minimum successful records not met: {len(records)} < {minimum}")
        context.log("INFO", "Crawl Links завершён", found=len(all_candidates), extracted=len(records), errors=len(errors))
        return {"records": records, "count": len(records), "discovered": len(all_candidates), "errors": errors,
                "failures": errors, "completed_urls": completed_urls, "resume_token": next_resume_token,
                "listing_diagnostics": listing_diagnostics, "detail_diagnostics": detail_diagnostics,
                "artifacts": list(context.artifacts)}

    @staticmethod
    def _detail_uses_browser(context: ExecutionContext, config: dict[str, Any]) -> bool:
        mode = str(config.get("detail_fetch_mode") or "AUTO").upper()
        if mode == "PLAYWRIGHT":
            return True
        if mode == "HTTP":
            return False
        source = context.variables.get("source", {})
        settings = source.get("settings", {}) if isinstance(source, dict) else {}
        profile = settings.get("profile", {}) if isinstance(settings, dict) else {}
        return str(source.get("fetch_mode") or settings.get("fetch_mode") or "").upper() == "PLAYWRIGHT" or bool(profile.get("requires_javascript"))

    def _effective_config(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        source = context.variables.get("source", {})
        settings = source.get("settings") if isinstance(source, dict) else {}
        settings = settings if isinstance(settings, dict) else {}
        source_crawl = settings.get("crawl_links") or settings.get("crawl") or {}
        source_crawl = source_crawl if isinstance(source_crawl, dict) else {}
        profile = settings.get("profile") if isinstance(settings.get("profile"), dict) else {}
        profile_extractor = profile.get("extractor") if isinstance(profile.get("extractor"), dict) else {}
        effective = dict(source_crawl)
        for key, value in config.items():
            if value not in (None, "", {}, []):
                effective[key] = value
        if not effective.get("link_selector"):
            fields = profile_extractor.get("fields") if isinstance(profile_extractor.get("fields"), list) else []
            link = next((field for field in fields if isinstance(field, dict) and field.get("attribute") == "href"), None)
            container = str(profile_extractor.get("container_selector") or "").strip()
            if link and link.get("selector"):
                effective["link_selector"] = f"{container} {link['selector']}".strip() if container else link["selector"]
        effective.setdefault("listing_url", str(source.get("url") or ""))
        effective.setdefault("same_origin_only", True)
        # Detail requests are rendered per candidate after URL regex groups
        # (for example publication timestamp and record id) become available.
        detail_request = effective.get("detail_request")
        rendered = render_object(effective, context, inputs)
        if isinstance(detail_request, dict):
            rendered["detail_request"] = detail_request
        return rendered

    async def _load_listing(
        self,
        context: ExecutionContext,
        inputs: dict[str, Any],
        config: dict[str, Any],
        client: httpx.AsyncClient | None = None,
    ) -> tuple[Any, str]:
        listing_url = render_template(str(config.get("listing_url") or ""), context, inputs)
        if not listing_url:
            return find_value(inputs, str(config.get("input_path") or "records")) or inputs, str(find_value(inputs, "url") or "")
        params = render_object(config.get("listing_query") or {}, context, inputs)
        date_range = config.get("date_range_query") or {}
        if isinstance(date_range, dict) and date_range:
            from_param = str(date_range.get("from_param") or "")
            to_param = str(date_range.get("to_param") or "")
            if not from_param or not to_param:
                raise ValueError("date_range_query requires from_param and to_param")
            try:
                zone = ZoneInfo(str(date_range.get("timezone") or "UTC"))
            except Exception as exc:
                raise ValueError("date_range_query contains an unknown timezone") from exc
            now = (context.effective_run_clock or datetime.now(UTC)).astimezone(zone)
            lookback_days = max(int(date_range.get("lookback_days") or 0), 0)
            pattern = str(date_range.get("format") or "YYYY-MM-DD")
            params.update({
                from_param: formula_format_date(now - timedelta(days=lookback_days), pattern),
                to_param: formula_format_date(now, pattern),
            })
        headers = {"Accept-Language": "ru-RU,ru;q=0.9,en;q=0.5", "User-Agent": "Mozilla/5.0 (compatible; ParserStudio/1.0)"}
        headers.update(render_object(config.get("headers") or {}, context, inputs))
        source = context.variables.get("source", {})
        source_settings = source.get("settings") if isinstance(source, dict) else {}
        source_settings = source_settings if isinstance(source_settings, dict) else {}
        profile = source_settings.get("profile") if isinstance(source_settings.get("profile"), dict) else {}
        source_fetch_mode = str(source.get("fetch_mode") or source_settings.get("fetch_mode") or "").upper() if isinstance(source, dict) else ""
        listing_fetch_mode = str(config.get("listing_fetch_mode") or "").upper()
        listing_uses_browser = listing_fetch_mode == "PLAYWRIGHT" or (
            listing_fetch_mode in {"", "AUTO"}
            and (source_fetch_mode == "PLAYWRIGHT" or profile.get("requires_javascript"))
        )
        if listing_uses_browser:
            rendered = await BrowserOpenNode().execute(
                context,
                inputs,
                {
                    "url": listing_url,
                    "wait_until": config.get("listing_wait_until", "networkidle"),
                    "timeout": config.get("listing_timeout", 60),
                    "headers": headers,
                    "capture_network": True,
                    "pagination_enabled": config.get("pagination_enabled", True),
                    "pagination_max_pages": config.get("pagination_max_pages", 25),
                    "pagination_next_selector": config.get("pagination_next_selector", ""),
                    "pagination_wait_ms": config.get("pagination_wait_ms", 500),
                    "tabs_enabled": config.get("tabs_enabled", False),
                    "tabs_wait_ms": config.get("tabs_wait_ms", 500),
                    "tabs_max_depth": config.get("tabs_max_depth", 4),
                    "full_page": False,
                    "http_fallback": True,
                    "egress_resolver": config.get("egress_resolver"),
                    "allowed_domains": config.get("allowed_domains"),
                    "egress_allowed_domains": config.get("egress_allowed_domains"),
                    "allowed_ports": config.get("allowed_ports"),
                    "egress_allowed_ports": config.get("egress_allowed_ports"),
                    "max_redirects": config.get("max_redirects"),
                },
            )
            # Merge every discovered JSON/XHR collection.  Taking only the
            # largest collection silently drops sibling tabs/categories on
            # sites that expose one endpoint per tab.
            network_payloads = [item.get("body") for item in rendered.get("network", []) if isinstance(item, dict) and item.get("body") is not None]
            if not config.get("items_path"):
                network_items = self._merge_url_lists(network_payloads)
                if network_items:
                    context.variables["_crawl_listing_diagnostics"] = {
                        "fetch_mode": "PLAYWRIGHT", "tab_count": rendered.get("tab_count", 0),
                        "tab_labels": rendered.get("tab_labels", []), "collection_source": "network_merged",
                    }
                    return network_items, str(rendered.get("url") or listing_url)
            context.variables["_crawl_listing_diagnostics"] = {
                "fetch_mode": "PLAYWRIGHT", "tab_count": rendered.get("tab_count", 0),
                "tab_labels": rendered.get("tab_labels", []), "collection_source": "rendered_html",
            }
            return rendered.get("html") or rendered.get("body") or "", str(rendered.get("url") or listing_url)
        listing_policy = FetchPolicy.from_config({**config, "request_timeout": config.get("listing_timeout", 60)})
        owns_client = client is None
        if client is None:
            cookies = render_object(config.get("cookies") or {}, context, inputs)
            client = httpx.AsyncClient(follow_redirects=False, headers=headers, cookies=cookies)
        responses: list[httpx.Response] = []
        documents: list[str] = []
        try:
            egress_policy = EgressPolicy.from_config(config)
            resolver = config.get("egress_resolver") or default_resolver
            response = await request_with_egress_policy(
                client,
                "GET",
                listing_url,
                listing_policy,
                egress_policy=egress_policy,
                resolver=resolver,
                request_fn=request_with_policy,
                params=params,
            )
            response.raise_for_status()
            responses.append(response)
            content_type = response.headers.get("content-type", "")
            if "json" not in content_type:
                documents.append(response.text)
            if config.get("pagination_enabled") and documents:
                maximum_pages = min(max(int(config.get("pagination_max_pages") or 25), 1), 500)
                selector = str(config.get("pagination_next_selector") or "a[rel='next']")
                seen_pages = {canonical_url(str(response.url))}
                while len(documents) < maximum_pages:
                    soup = BeautifulSoup(documents[-1], "lxml")
                    next_link = soup.select_one(selector) or soup.select_one("a[rel='next'][href]")
                    next_href = str(next_link.get("href") or "").strip() if next_link else ""
                    if not next_href:
                        break
                    next_url = canonical_url(urljoin(str(response.url), next_href))
                    if next_url in seen_pages:
                        break
                    seen_pages.add(next_url)
                    response = await request_with_egress_policy(
                        client,
                        "GET",
                        next_url,
                        listing_policy,
                        egress_policy=egress_policy,
                        resolver=resolver,
                        request_fn=request_with_policy,
                    )
                    response.raise_for_status()
                    responses.append(response)
                    documents.append(response.text)
        finally:
            if owns_client:
                await close_http_client(client)
        context.variables["_crawl_listing_diagnostics"] = {
            "fetch_mode": "HTTP",
            "fetch_attempts": [
                attempt
                for page_response in responses
                for attempt in fetch_attempts(page_response)
            ],
            "pages": len(responses),
            "redirect_chain": [
                hop
                for page_response in responses
                for hop in page_response.extensions.get("redirect_chain", [])
            ],
        }
        listing_body = merge_rendered_sections(documents) if len(documents) > 1 else response.text
        if config.get("save_artifacts", True):
            artifact_content = listing_body.encode("utf-8") if documents else response.content
            await store_artifact(context, artifact_content, response.headers.get("content-type", "application/octet-stream"), str(response.url), "listing.json" if "json" in response.headers.get("content-type", "") else "listing.html", "raw_listing")
        if "json" in response.headers.get("content-type", ""):
            return response.json(), str(response.url)
        return listing_body, str(responses[0].url)

    def _listing_items(self, listing: Any, config: dict[str, Any]) -> list[Any]:
        if isinstance(listing, str):
            soup = BeautifulSoup(listing, "lxml")
            selector = str(config.get("link_selector") or "").strip()
            elements = soup.select(selector) if selector else soup.select("main a[href]") or soup.select("a[href]")
            items: list[dict[str, Any]] = []
            for element in elements:
                href = str(element.get("href") or "").strip()
                if not href or href.startswith("#") or href.lower().startswith(("javascript:", "mailto:", "tel:")):
                    continue
                utility_path = urlsplit(href).path.rstrip("/").lower() or "/"
                if not selector and utility_path in {"/", "/search"}:
                    continue
                if not selector and any(
                    parent.name in {"header", "footer", "nav"}
                    or "breadcrumb" in " ".join(parent.get("class") or []).lower()
                    or "breadcrumb" in str(parent.get("id") or "").lower()
                    for parent in element.parents
                ):
                    continue
                items.append({"url": href, "title": element.get_text(" ", strip=True)})
            return items
        selected = find_value(listing, str(config.get("items_path") or "")) if config.get("items_path") else self._largest_url_list(listing)
        return selected if isinstance(selected, list) else []

    def _largest_url_list(self, value: Any) -> list[Any]:
        candidates: list[list[Any]] = []
        def visit(item: Any) -> None:
            if isinstance(item, list):
                if item and sum(1 for child in item if isinstance(child, dict) and child.get("url")):
                    candidates.append(item)
                for child in item:
                    visit(child)
            elif isinstance(item, dict):
                for child in item.values():
                    visit(child)
        visit(value)
        return max(candidates, key=lambda items: (sum(1 for child in items if isinstance(child, dict) and child.get("url")), len(items)), default=[])

    def _merge_url_lists(self, value: Any) -> list[Any]:
        """Flatten and deduplicate all nested URL-bearing collections."""
        merged: list[Any] = []
        seen: set[str] = set()
        for items in self._url_lists(value):
            for item in items:
                if isinstance(item, dict):
                    raw = item.get("url") or item.get("href") or item.get("link")
                    if not raw:
                        continue
                    key = str(raw)
                    if key in seen:
                        continue
                    seen.add(key)
                    if "url" not in item:
                        item = {**item, "url": raw}
                    merged.append(item)
                elif isinstance(item, str) and item.startswith(("http://", "https://", "/")):
                    if item not in seen:
                        seen.add(item)
                        merged.append({"url": item})
        return merged

    def _url_lists(self, value: Any) -> list[list[Any]]:
        found: list[list[Any]] = []
        def visit(item: Any) -> None:
            if isinstance(item, list):
                url_count = sum(1 for child in item if isinstance(child, dict) and (child.get("url") or child.get("href") or child.get("link")))
                if url_count:
                    found.append(item)
                for child in item:
                    visit(child)
            elif isinstance(item, dict):
                for child in item.values():
                    visit(child)
        visit(value)
        return found


class ParseHTMLNode:
    type = "parse_html"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        html = find_value(inputs, str(config.get("input_path", "body")))
        if not isinstance(html, str):
            raise ValueError("Parse HTML ожидает строку HTML")
        soup = BeautifulSoup(html, "lxml")
        tables = []
        for table in soup.select("table"):
            tables.append([[cell.get_text(" ", strip=True) for cell in row.select("th,td")] for row in table.select("tr")])
        return {"html": html, "text": soup.get_text("\n", strip=True), "links": [a.get("href") for a in soup.select("a[href]")], "tables": tables, "metadata": {"title": soup.title.get_text(strip=True) if soup.title else "", "language": soup.html.get("lang") if soup.html else None}}


class SelectElementsNode:
    type = "select_elements"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        html = find_value(inputs, str(config.get("input_path", "html"))) or find_value(inputs, "body")
        attr = config.get("attribute")
        mode = config.get("mode", "text")
        values: list[Any] = []
        evidence: list[dict[str, Any]] = []
        if config.get("xpath"):
            document = lxml_html.fromstring(str(html))
            for element in document.xpath(str(config["xpath"])):
                value = element.get(attr) if attr and hasattr(element, "get") else etree.tostring(element, encoding="unicode") if mode == "html" else " ".join(element.itertext()).strip() if hasattr(element, "itertext") else str(element)
                values.append(value)
                evidence.append({"xpath": config["xpath"], "text": str(value)[:500]})
        else:
            selector = config.get("selector")
            if not selector:
                raise ValueError("CSS selector или XPath не задан")
            soup = BeautifulSoup(str(html), "lxml")
            for element in soup.select(str(selector)):
                value = element.get(attr) if attr else str(element) if mode == "html" else element.get_text(" ", strip=True)
                values.append(value)
                evidence.append({"css_selector": selector, "text": str(value)[:500]})
        if config.get("single"):
            return {"value": values[0] if values else None, "count": len(values), "evidence": evidence[:1]}
        return {"items": values, "records": values, "count": len(values), "evidence": evidence}


class ExtractRepeatingListNode:
    type = "extract_repeating_list"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        html = find_value(inputs, str(config.get("input_path", "html"))) or find_value(inputs, "body")
        selector = str(config.get("container_selector") or "")
        if not selector:
            raise ValueError("Selector карточки не задан")
        soup = BeautifulSoup(str(html), "lxml")
        rows: list[dict[str, Any]] = []
        for container in soup.select(selector):
            row: dict[str, Any] = {}
            evidence: dict[str, Any] = {}
            for field in config.get("fields", []):
                name = field.get("name")
                if not name:
                    continue
                # Field selectors are relative to a repeating container.  A
                # card is often the link itself, so the no-code editor needs
                # an explicit way to select that container rather than an
                # impossible descendant ``a.card``.  ``:scope`` (and the
                # concise ``.`` alias) mean the current container.
                field_selector = str(field.get("selector") or "").strip()
                element = container if field_selector in {":scope", "."} else container.select_one(field_selector)
                if not element:
                    row[name] = field.get("default")
                    continue
                if field.get("attribute"):
                    row[name] = element.get(field["attribute"])
                elif field.get("mode") == "html":
                    row[name] = str(element)
                else:
                    row[name] = element.get_text(" ", strip=True)
                    # A few public RSS generators emit ``<link/>`` followed
                    # by the URL as the element's XML tail instead of putting
                    # it inside ``<link>``.  BeautifulSoup models that URL as
                    # a sibling string, so ordinary CSS text extraction sees
                    # an empty link.  Treat the immediately following
                    # non-empty text node as the link value only for an empty
                    # XML/RSS ``link`` element.  This is content-shape
                    # handling, not a site-specific preset, and normal Atom
                    # links still use their configured ``href`` attribute.
                    if (
                        not row[name]
                        and element.name == "link"
                        and "xml" in str(html).lstrip()[:200].lower()
                    ):
                        sibling = element.next_sibling
                        while sibling is not None and not str(sibling).strip():
                            sibling = sibling.next_sibling
                        candidate = str(sibling).strip() if sibling is not None else ""
                        if candidate.startswith(("http://", "https://")):
                            row[name] = candidate
                evidence[name] = {"css_selector": selector if element is container else f"{selector} {field_selector}", "text": str(row[name])[:500]}
            row.setdefault("evidence", evidence)
            rows.append(row)
        return {"records": rows, "count": len(rows)}


class ParseTableNode:
    type = "parse_table"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        html = find_value(inputs, str(config.get("input_path", "html"))) or find_value(inputs, "body")
        soup = BeautifulSoup(str(html), "lxml")
        selector = str(config.get("selector", "table"))
        matches = soup.select(selector)
        if not matches:
            raise ValueError("Таблица не найдена")
        # ``table_index`` picks which occurrence of the selector to parse; the
        # chosen occurrence becomes part of the row identity so two tables on
        # one page can never collapse into the same natural key.
        table_index = max(int(config.get("table_index", 0) or 0), 0)
        if table_index >= len(matches):
            table_index = 0
        table = matches[table_index]
        table_id = f"{selector}:{table_index}"
        matrix: list[list[str]] = []
        rowspans: dict[int, tuple[str, int]] = {}
        for tr in table.select("tr"):
            row: list[str] = []
            col = 0
            cells = tr.select(":scope > th, :scope > td")
            for cell in cells:
                while col in rowspans and rowspans[col][1] > 0:
                    value, remaining = rowspans[col]
                    row.append(value)
                    rowspans[col] = (value, remaining - 1)
                    col += 1
                value = cell.get_text(" ", strip=True)
                colspan = max(int(cell.get("colspan", 1)), 1)
                rowspan = max(int(cell.get("rowspan", 1)), 1)
                for _ in range(colspan):
                    row.append(value)
                    if rowspan > 1:
                        rowspans[col] = (value, rowspan - 1)
                    col += 1
            while col in rowspans and rowspans[col][1] > 0:
                value, remaining = rowspans[col]
                row.append(value)
                rowspans[col] = (value, remaining - 1)
                col += 1
            if any(row):
                matrix.append(row)
        header_row = int(config.get("header_row", 0))
        headers = dedupe_headers(matrix[header_row]) if matrix else []
        records = [dict(zip(headers, row, strict=False)) for row in matrix[header_row + 1:] if any(row)]
        if config.get("normalize_fields"):
            for record in records:
                for header, value in list(record.items()):
                    normalized = normalize_table_field_name(header)
                    if normalized and normalized not in record:
                        record[normalized] = value
        # Every row carries a structural identity (page/table/row) so distinct
        # rows survive the Output natural-key check even when no business
        # identity column is mapped yet.  ``setdefault`` keeps a real column
        # that happens to share the name.
        page_url = _resolved_page_url(inputs)
        for row_index, record in enumerate(records):
            record.setdefault("row_index", row_index)
            record.setdefault("table_id", table_id)
            if page_url:
                record.setdefault("page_url", page_url)
        columns = [
            {"index": index, "header": header, "sample": str(records[0].get(header, ""))[:120] if records else ""}
            for index, header in enumerate(headers)
        ]
        mapping_draft = [
            {"header": header, "field": normalize_table_field_name(header) or f"column_{index}"}
            for index, header in enumerate(headers)
            if str(header).strip()
        ]
        return {
            "records": records,
            "count": len(records),
            "headers": headers,
            "table": matrix,
            "columns": columns,
            "mapping_draft": mapping_draft,
        }


def _resolved_page_url(inputs: dict[str, Any]) -> str:
    bundle = inputs.get("source_bundle") if isinstance(inputs.get("source_bundle"), Mapping) else {}
    value = inputs.get("url") or bundle.get("final_url") or inputs.get("final_url") or bundle.get("seed_url")
    return str(value) if value else ""


_DATE_LIKE_RE = re.compile(r"\b\d{2}[./-]\d{2}[./-]\d{2,4}\b|\b\d{4}-\d{2}-\d{2}\b")
_RATE_LIKE_RE = re.compile(r"\d+[.,]?\d*\s*%|\b\d+[.,]\d{2,}\b|\b(?:BYN|RUB|USD|EUR|PLN|BYR)\b", re.I)
_CLASS_TOKEN_RE = re.compile(r"^[A-Za-z_][\w-]*$")


def _card_selector(element: Any) -> str | None:
    """Build a stable CSS signature for a repeating container candidate."""

    classes = sorted({str(item) for item in (element.get("class") or []) if _CLASS_TOKEN_RE.match(str(item))})
    if not classes:
        return None
    return f"{element.name}." + ".".join(classes)


def detect_card_clusters(html: Any, *, limit: int = 8) -> list[dict[str, Any]]:
    """Detect repeating card containers in any public markup.

    Scoring is purely structural: candidate containers are grouped by a
    tag+class signature, then ranked by text density and link/price/date
    signal density.  Navigation menus and one-off blocks lose to real card
    grids because their per-item text is short and signal-free.  No domain,
    URL or site-specific branch is involved.
    """

    soup = BeautifulSoup(str(html), "lxml")
    groups: dict[str, list[Any]] = {}
    for element in soup.find_all(True):
        signature = _card_selector(element)
        if signature:
            groups.setdefault(signature, []).append(element)
    candidates: list[dict[str, Any]] = []
    for selector, items in groups.items():
        count = len(items)
        if not 3 <= count <= 500:
            continue
        probe = items[:20]
        texts = [item.get_text(" ", strip=True) for item in probe]
        avg_text = sum(len(text) for text in texts) / len(texts)
        if avg_text < 25:
            continue
        link_fraction = sum(1 for item in probe if item.select_one("a[href]")) / len(probe)
        signal_fraction = sum(1 for text in texts if _DATE_LIKE_RE.search(text) or _RATE_LIKE_RE.search(text)) / len(texts)
        descendant_count = sum(len(item.find_all()) for item in items[:5]) / min(count, 5)
        score = (round(signal_fraction, 3), round(link_fraction, 3), min(round(avg_text), 2000), count)
        candidates.append({
            "selector": selector,
            "count": count,
            "avg_text_length": round(avg_text, 1),
            "link_fraction": round(link_fraction, 3),
            "signal_fraction": round(signal_fraction, 3),
            "descendant_count": round(descendant_count, 1),
            "score": list(score),
            "sample_text": texts[0][:200],
        })
    candidates.sort(key=lambda item: tuple(item["score"]), reverse=True)
    return candidates[:limit]


def card_cluster_passes(candidate: Mapping[str, Any]) -> bool:
    """Minimal structural bar a cluster must clear to be auto-selected.

    Two equally legitimate public shapes pass: card grids built from links
    (classic listings, deposits, news), and link-less offer panels whose only
    hooks are price/term/date signals (button-driven offers, MUI-style
    catalog grids).  Navigation noise fails both bars: short, signal-free
    text.
    """

    avg_text = float(candidate.get("avg_text_length") or 0)
    link_fraction = float(candidate.get("link_fraction") or 0)
    signal_fraction = float(candidate.get("signal_fraction") or 0)
    return (
        (avg_text >= 40 and link_fraction >= 0.2)
        or (avg_text >= 60 and signal_fraction >= 0.3)
    )


class JSONPathNode:
    type = "json_path"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        data = find_value(inputs, str(config.get("input_path", "body")))
        if data is None:
            return {"items": [], "records": [], "count": 0}
        path = str(config.get("path", "$"))
        try:
            from jsonpath_ng.ext import parse
            values = [match.value for match in parse(path).find(data)]
        except Exception:
            values = simple_json_path(data, path)
        return {"items": values, "records": values, "count": len(values)}


class ParseDocumentNode:
    type = "parse_document"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        encoded = find_value(inputs, str(config.get("input_path", "content_base64")))
        filename = str(find_value(inputs, str(config.get("filename_path", "filename"))) or config.get("filename") or "document")
        if not encoded:
            raise ValueError("Parse Document: base64-содержимое не найдено")
        data = base64.b64decode(str(encoded))
        suffix = Path(filename).suffix.lower()
        if suffix == ".csv":
            text = data.decode("utf-8-sig")
            try:
                dialect = csv.Sniffer().sniff(text[:4096])
            except csv.Error:
                dialect = csv.excel
            records = list(csv.DictReader(io.StringIO(text), dialect=dialect))
            return {"type": "CSV", "records": records, "count": len(records), "evidence": {"filename": filename}}
        if suffix == ".xlsx":
            workbook = load_workbook(io.BytesIO(data), data_only=not bool(config.get("formulas")), read_only=True)
            sheets: dict[str, list[dict[str, Any]]] = {}
            selected = str(config.get("sheet") or "")
            for sheet in workbook.worksheets:
                if selected and sheet.title != selected:
                    continue
                rows = list(sheet.iter_rows(values_only=True))
                header_row = int(config.get("header_row", 0))
                headers = dedupe_headers([str(value) if value is not None else "" for value in rows[header_row]]) if len(rows) > header_row else []
                sheets[sheet.title] = [dict(zip(headers, row, strict=False)) for row in rows[header_row + 1:] if any(value is not None for value in row)]
            records = next(iter(sheets.values()), [])
            return {"type": "XLSX", "sheets": sheets, "records": records, "count": len(records), "evidence": {"filename": filename}}
        if suffix == ".docx":
            from docx import Document
            document = Document(io.BytesIO(data))
            tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
            return {"type": "DOCX", "paragraphs": [p.text for p in document.paragraphs], "tables": tables, "text": "\n".join(p.text for p in document.paragraphs), "evidence": {"filename": filename}}
        if suffix == ".pdf":
            if config.get("use_docling", True):
                try:
                    import tempfile

                    from docling.document_converter import DocumentConverter
                    with tempfile.NamedTemporaryFile(suffix=".pdf") as temporary:
                        temporary.write(data)
                        temporary.flush()
                        conversion = DocumentConverter().convert(temporary.name)
                    markdown = conversion.document.export_to_markdown()
                    document_json = conversion.document.export_to_dict()
                    return {"type": "PDF", "parser": "DOCLING", "text": markdown, "markdown": markdown, "document": document_json, "records": document_json.get("tables", []), "evidence": {"filename": filename}}
                except Exception as exc:
                    context.log("WARNING", "Docling недоступен; использован pypdf", error=str(exc))
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            selected_pages = parse_page_selection(str(config.get("pages") or ""), len(reader.pages))
            pages = [{"page": index + 1, "text": reader.pages[index].extract_text() or ""} for index in selected_pages]
            return {"type": "PDF", "parser": "PYPDF", "pages": pages, "page_count": len(pages), "text": "\n".join(page["text"] for page in pages), "evidence": {"filename": filename}}
        if suffix == ".json":
            parsed = json.loads(data)
            return {"type": "JSON", "data": parsed, "records": parsed if isinstance(parsed, list) else [parsed]}
        raise ValueError(f"Неподдерживаемый документ: {suffix}")


class TransformNode:
    type = "transform"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        data = find_value(inputs, str(config.get("input_path", "records")))
        if data is None:
            data = inputs
        records = data if isinstance(data, list) else [data]
        output: list[dict[str, Any]] = []
        decisions: list[dict[str, Any]] = []
        transformed = [dict(item) if isinstance(item, dict) else {"value": item} for item in records]
        for operation in config.get("operations", []):
            if str(operation.get("type") or "") in _COLLECTION_OPERATION_TYPES:
                transformed = apply_collection_operation(transformed, operation, context)
            else:
                for row in transformed:
                    apply_operation(row, operation, context=context)
        for row in transformed:
            original = dict(row)
            include, reason = deterministic_filter_decision(row, config.get("filters") or [])
            if include:
                output.append(row)
                decisions.append({"action": "INCLUDED", "reason": reason, "record": _decision_identity(row)})
            else:
                decisions.append({"action": "EXCLUDED", "reason": reason, "record": _decision_identity(original)})
        identity_fields = config.get("identity") or config.get("identity_fields") or []
        if isinstance(identity_fields, str):
            identity_fields = [value.strip() for value in identity_fields.split(",") if value.strip()]
        deduped, duplicates = deterministic_deduplicate(output, identity_fields)
        decisions.extend(duplicates)
        transform_revision = str(config.get("transformRevision") or config.get("transform_revision") or "process@2")
        # Keep pass-through records byte/semantically compatible.  Process
        # provenance starts when the preset actually declares a transformation
        # or identity/filter decision.
        if config.get("operations") or config.get("filters") or identity_fields:
            for record in deduped:
                provenance = record.get("__provenance") if isinstance(record.get("__provenance"), dict) else {}
                record["__provenance"] = {
                    **provenance,
                    "process": {"transform_revision": transform_revision},
                }
        return {
            "records": deduped,
            "count": len(deduped),
            "business_records": bool(inputs.get("business_records")),
            "filter_decisions": decisions,
            "duplicates_removed": len(duplicates),
            "transform_revision": transform_revision,
        }


class MappingNode:
    """Explicitly turns a transport envelope into dataset business records."""

    type = "mapping"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        if config.get("fieldCandidates") or config.get("field_candidates"):
            return schema_first_extract(context, inputs, config)
        value = find_value(inputs, str(config.get("input_path", "records")))
        rows = value if isinstance(value, list) else [value] if value is not None else []
        records: list[dict[str, Any]] = []
        errors: list[dict[str, Any]] = []
        for index, item in enumerate(rows):
            source = item if isinstance(item, dict) else {"value": item}
            record: dict[str, Any] = {}
            if not config.get("fields"):
                records.append(dict(source))
                continue
            for spec in config.get("fields", []):
                target = str(spec.get("target") or spec.get("name") or "")
                if not target:
                    errors.append({"row": index, "code": "TARGET_REQUIRED"})
                    continue
                if "constant" in spec and spec.get("constant") is not None:
                    result = spec["constant"]
                elif spec.get("expression"):
                    try:
                        result = safe_eval(str(spec["expression"]), source, context.effective_run_clock)
                    except Exception as exc:
                        errors.append({"row": index, "field": target, "code": "EXPRESSION", "message": str(exc)})
                        result = spec.get("default")
                else:
                    result = find_value(source, str(spec.get("source_path") or spec.get("source") or target))
                    if result is None:
                        result = spec.get("default")
                if spec.get("required") and result in (None, ""):
                    errors.append({"row": index, "field": target, "code": "REQUIRED"})
                record[target] = result
                evidence = source.get("evidence") if isinstance(source.get("evidence"), dict) else {}
                if target in evidence:
                    record.setdefault("__evidence", {})[target] = evidence[target]
            provenance = source.get("__provenance")
            if isinstance(provenance, dict):
                record["__provenance"] = provenance
            records.append(record)
        return {"records": records, "count": len(records), "mapping_errors": errors,
                "business_records": True, "schema_preview": records[:5]}


def schema_first_extract(context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    """Extract a target schema from ordered field candidates.

    Candidates are declarative JSON/JSON-LD/DOM/table/document paths.  A field
    is accepted only after its type/cardinality postconditions pass, and the
    chosen candidate remains attached to the output record as evidence.
    """

    candidates_by_field = config.get("fieldCandidates") or config.get("field_candidates") or {}
    if not isinstance(candidates_by_field, dict):
        raise ValueError("fieldCandidates must be an object keyed by target field")
    collection = _schema_collection(inputs, config)
    selection = str(config.get("selection") or "first_valid").lower()
    mode = str(config.get("mode") or "AUTO").upper()
    records: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    for index, source in enumerate(collection):
        row: dict[str, Any] = {}
        evidence: dict[str, Any] = {}
        for target, raw_candidates in candidates_by_field.items():
            candidates = raw_candidates if isinstance(raw_candidates, list) else []
            choices: list[dict[str, Any]] = []
            for position, candidate in enumerate(candidates):
                if not isinstance(candidate, dict):
                    continue
                value, pointer = _extract_field_candidate(source, candidate, inputs)
                result = _candidate_result(target, value, candidate, pointer, position)
                choices.append(result)
            accepted = [item for item in choices if item["passed"]]
            pinned = str(config.get("selectedCandidates", {}).get(target, "")) if isinstance(config.get("selectedCandidates"), dict) else ""
            if mode in {"MANUAL", "ASSISTED"} and pinned:
                accepted = [item for item in accepted if item["id"] == pinned]
            chosen: dict[str, Any] | None = None
            if selection == "best_coverage":
                chosen = max(accepted, key=lambda item: (item["coverage"], -item["position"]), default=None)
            elif selection == "merge_non_conflicting":
                values = {json.dumps(item["value"], sort_keys=True, default=str) for item in accepted}
                chosen = accepted[0] if len(values) == 1 and accepted else None
                if len(values) > 1:
                    errors.append({"row": index, "field": target, "code": "CANDIDATE_CONFLICT", "candidates": choices})
            else:
                chosen = accepted[0] if accepted else None
            if chosen is None:
                if any(bool(item.get("required")) for item in candidates if isinstance(item, dict)):
                    errors.append({"row": index, "field": target, "code": "FIELD_CANDIDATES_EXHAUSTED", "candidates": choices})
                row[target] = None
                continue
            row[target] = chosen["value"]
            evidence[target] = {
                "candidate": chosen["id"],
                "kind": chosen["kind"],
                "pointer": chosen["pointer"],
                "coverage": chosen["coverage"],
                "source_value": chosen["value"],
            }
        provenance = source.get("__provenance") if isinstance(source, dict) and isinstance(source.get("__provenance"), dict) else {}
        row["__provenance"] = {**provenance, "field_evidence": evidence}
        records.append(row)
    return {
        "records": records,
        "count": len(records),
        "mapping_errors": errors,
        "business_records": True,
        "schema_ref": config.get("targetSchemaRef") or config.get("target_schema_ref"),
        "field_evidence": [record.get("__provenance", {}).get("field_evidence", {}) for record in records],
    }


def _schema_collection(inputs: dict[str, Any], config: dict[str, Any]) -> list[dict[str, Any]]:
    configured = find_value(inputs, str(config.get("collectionPath") or config.get("collection_path") or "records"))
    if isinstance(configured, list):
        return [item if isinstance(item, dict) else {"value": item} for item in configured]
    if configured is not None:
        return [configured if isinstance(configured, dict) else {"value": configured}]
    return [inputs]


def _extract_field_candidate(source: dict[str, Any], candidate: dict[str, Any], inputs: dict[str, Any]) -> tuple[Any, str]:
    kind = str(candidate.get("kind") or candidate.get("source") or "path").lower()
    path = str(candidate.get("path") or candidate.get("source_path") or candidate.get("selector") or "")
    if kind in {"json", "path", "jsonpath", "listing", "document"}:
        root = source if kind != "document" else (source.get("document") or source)
        if path.startswith("$"):
            values = simple_json_path(root, path)
            return (values if candidate.get("multiple") else values[0] if values else None), path
        return find_value(root, path), path
    if kind in {"json_ld", "jsonld"}:
        html = str(source.get("html") or source.get("body") or inputs.get("html") or inputs.get("body") or "")
        values: list[Any] = []
        for script in BeautifulSoup(html, "lxml").select("script[type='application/ld+json']"):
            try:
                payload = json.loads(script.get_text())
            except (TypeError, json.JSONDecodeError):
                continue
            extracted = simple_json_path(payload, path if path.startswith("$") else f"$.{path}")
            values.extend(extracted)
        return (values if candidate.get("multiple") else values[0] if values else None), path
    if kind in {"dom", "css", "table"}:
        html = str(source.get("html") or source.get("body") or inputs.get("html") or inputs.get("body") or "")
        soup = BeautifulSoup(html, "lxml")
        elements = soup.select(path)
        attribute = candidate.get("attribute")
        values = [element.get(attribute) if attribute else element.get_text(" ", strip=True) for element in elements]
        return (values if candidate.get("multiple") else values[0] if values else None), path
    if kind == "constant":
        return candidate.get("value"), "constant"
    return None, path


def _candidate_result(target: str, value: Any, candidate: dict[str, Any], pointer: str, position: int) -> dict[str, Any]:
    multiple = bool(candidate.get("multiple"))
    required = bool(candidate.get("required", True))
    nonempty = bool(value) if multiple else value not in (None, "")
    expected = str(candidate.get("valueType") or candidate.get("value_type") or "")
    type_ok = _value_matches_type(value, expected, multiple)
    coverage = 1.0 if nonempty else 0.0
    minimum = float(candidate.get("minCoverage", candidate.get("min_coverage", 1.0 if required else 0.0)) or 0.0)
    return {
        "id": str(candidate.get("id") or f"{target}:{position}"),
        "position": position,
        "kind": str(candidate.get("kind") or candidate.get("source") or "path"),
        "pointer": pointer,
        "value": value,
        "required": required,
        "coverage": coverage,
        "passed": nonempty and type_ok and coverage >= minimum,
    }


def _value_matches_type(value: Any, expected: str, multiple: bool) -> bool:
    if value in (None, ""):
        return False
    if multiple:
        return isinstance(value, list)
    if not expected:
        return True
    if expected in {"string", "text"}:
        return isinstance(value, str)
    if expected in {"number", "decimal"}:
        return isinstance(value, (int, float)) or normalize_number(value) is not None
    if expected in {"object", "json"}:
        return isinstance(value, dict)
    if expected == "array":
        return isinstance(value, list)
    return True


def deterministic_filter_decision(row: dict[str, Any], rules: list[Any]) -> tuple[bool, str]:
    for raw in rules:
        if not isinstance(raw, dict):
            continue
        field = str(raw.get("field") or "")
        value = find_value(row, field)
        operation = str(raw.get("operator") or "equals").lower()
        expected = raw.get("value")
        matched = {
            "equals": value == expected,
            "not_equals": value != expected,
            "exists": value not in (None, ""),
            "empty": value in (None, ""),
            "contains": str(expected) in str(value or ""),
            "regex": bool(re.search(str(expected or ""), str(value or ""), re.I)),
            "in": value in (expected if isinstance(expected, list) else [expected]),
        }.get(operation, False)
        if matched and str(raw.get("action") or "exclude").lower() == "exclude":
            return False, str(raw.get("reason") or f"{field}:{operation}")
        if not matched and str(raw.get("action") or "").lower() == "include_only":
            return False, str(raw.get("reason") or f"{field}:{operation}:not_matched")
    return True, "PASSED_DETERMINISTIC_RULES"


def deterministic_deduplicate(records: list[dict[str, Any]], keys: list[Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    seen: set[str] = set()
    unique: list[dict[str, Any]] = []
    decisions: list[dict[str, Any]] = []
    for row in records:
        identity = [find_value(row, str(key)) for key in keys] if keys else row
        key = json.dumps(identity, ensure_ascii=False, sort_keys=True, default=str)
        if key in seen:
            decisions.append({"action": "DEDUPLICATED", "reason": "NATURAL_IDENTITY", "record": _decision_identity(row)})
            continue
        seen.add(key)
        unique.append(row)
    return unique, decisions


def _decision_identity(row: dict[str, Any]) -> str:
    return hashlib.sha256(json.dumps(row, ensure_ascii=False, sort_keys=True, default=str).encode()).hexdigest()[:16]


class SetConstantNode:
    """Produces a record or record array without requiring a hand-written graph."""
    type = "set_constant"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        value = config.get("value", config.get("object", {}))
        value = render_object(value, context, inputs)
        records = value if isinstance(value, list) else [value]
        records = [item if isinstance(item, dict) else {"value": item} for item in records]
        return {"records": records, "count": len(records), "business_records": True}


class FormulaNode:
    type = "formula"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        data = find_value(inputs, str(config.get("input_path", "records")))
        records = data if isinstance(data, list) else [data]
        expression = str(config.get("expression") or "")
        target = str(config.get("target") or "value")
        output = []
        for item in records:
            row = dict(item) if isinstance(item, dict) else {"value": item}
            row[target] = safe_eval(expression, row, context.effective_run_clock)
            output.append(row)
        return {"records": output, "count": len(output)}


class LLMExtractNode:
    type = "llm_extract"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        content = find_value(inputs, str(config.get("input_path", "text")))
        provider = str(config.get("provider", "deepseek"))
        provider_config = context.capabilities.get("ai_providers", {}).get(provider, {})
        base_url = str(provider_config.get("base_url") or context.variables.get("deepseek_base_url") or "https://api.deepseek.com").rstrip("/")
        api_key = str(provider_config.get("api_key") or context.secrets.get(f"AI_PROVIDER_{provider}") or context.secrets.get("DEEPSEEK_API_KEY") or "")
        model = str(config.get("model") or provider_config.get("default_model") or "deepseek-chat")
        system_prompt = render_template(str(config.get("system_prompt") or "Верни только валидный JSON."), context, inputs)
        if "максимум 50" not in system_prompt.lower():
            system_prompt += "\nВерни максимум 50 записей за один запуск; при избытке выбери только наиболее актуальные и полные."
        schema = config.get("response_schema") or {}
        user_template = str(config.get("user_prompt") or "Извлеки данные из:\n{{content}}")
        max_input_chars = max(int(config.get("max_input_chars", 36000)), 1000)
        prompt_content = content[:max_input_chars] if isinstance(content, str) else content
        user_prompt = user_template.replace("{{content}}", stringify(prompt_content)).replace("{{schema}}", json.dumps(schema, ensure_ascii=False))
        if provider == "mock":
            parsed = config.get("mock_response") or ({"records": content} if isinstance(content, list) else {"value": content})
            return llm_output(parsed, "mock", {"prompt_tokens": 0, "completion_tokens": 0})
        if not api_key:
            if config.get("fallback_to_input"):
                return {"records": content if isinstance(content, list) else [content], "llm_skipped": True, "reason": "API key отсутствует"}
            raise ValueError(f"API key для provider {provider} не настроен")
        payload: dict[str, Any] = {
            "model": model,
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            "temperature": float(config.get("temperature", 0)), "max_tokens": int(config.get("max_tokens", 3000)),
        }
        if config.get("json_mode", True):
            payload["response_format"] = {"type": "json_object"}
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        try:
            async with httpx.AsyncClient(timeout=float(config.get("timeout", 60))) as client:
                response = await client.post(f"{base_url}/chat/completions", headers=headers, json=payload)
                response.raise_for_status()
            raw = response.json()
            text = raw["choices"][0]["message"]["content"]
            parsed = parse_json_response(text)
            # JSON-mode providers commonly wrap an array in a named envelope
            # (for example {"products": [...]}) despite an array response
            # contract. Treat that as the requested collection, not a failure.
            if isinstance(parsed, dict):
                for key in ("records", "items", "products", "data", "results"):
                    if isinstance(parsed.get(key), list) and (schema.get("type") == "array" or len(parsed) == 1):
                        parsed = parsed[key]
                        break
            if isinstance(parsed, list):
                parsed = dedupe_extracted_records(parsed, config.get("dedupe_key_fields"))
            validate_json_schema(parsed, schema)
            return {**llm_output(parsed, raw.get("model", model), raw.get("usage", {})), "response": text}
        except Exception:
            if config.get("fallback_to_input"):
                return {"records": content if isinstance(content, list) else [content], "llm_fallback": True}
            raise

def dedupe_extracted_records(records: list[Any], key_fields: list[str] | None = None) -> list[Any]:
    """Discard repeated cards emitted by templated pages or an LLM."""
    unique: list[Any] = []
    seen: set[str] = set()
    for record in records:
        if not isinstance(record, dict):
            unique.append(record)
            continue
        key = "|".join(str(record.get(field) or "") for field in key_fields) if key_fields else json.dumps(record, ensure_ascii=False, sort_keys=True, default=str)
        if key in seen:
            continue
        seen.add(key)
        unique.append(record)
    return unique


class LLMClassifyNode:
    type = "llm_classify"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        labels = config.get("labels") or []
        extract_config = {**config, "response_schema": {"type": "object", "required": ["label"], "properties": {"label": {"type": "string", "enum": labels}}}, "system_prompt": "Классифицируй значение. Верни JSON вида {\"label\": ...}."}
        result = await LLMExtractNode().execute(context, inputs, extract_config)
        parsed = result.get("parsed_response") or {}
        return {**result, "value": parsed.get("label")}


class ValidateNode:
    type = "validate"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        records = find_value(inputs, str(config.get("input_path", "records")))
        records = records if isinstance(records, list) else [records]
        errors: list[dict[str, Any]] = []
        quarantined: list[dict[str, Any]] = []
        for index, row in enumerate(records):
            for field in config.get("required", []):
                if not isinstance(row, dict) or row.get(field) in (None, ""):
                    errors.append({"row": index, "field": field, "code": "REQUIRED"})
            for rule in config.get("ranges", []):
                value = row.get(rule["field"]) if isinstance(row, dict) else None
                if value is not None and (value < rule.get("min", value) or value > rule.get("max", value)):
                    errors.append({"row": index, "field": rule["field"], "code": "RANGE"})
            try:
                validate_json_schema(row, config.get("schema") or {})
            except ValueError as exc:
                errors.append({"row": index, "code": "SCHEMA", "message": str(exc)})
        traversal = inputs.get("traversal") if isinstance(inputs.get("traversal"), dict) else {}
        reconciliation = traversal.get("reconciliation") if isinstance(traversal.get("reconciliation"), dict) else {}
        reconciliation = {
            "discovered": int(reconciliation.get("discovered", len(records)) or 0),
            "succeeded": int(reconciliation.get("succeeded", len(records)) or 0),
            "intentionally_skipped": int(reconciliation.get("intentionally_skipped", 0) or 0),
            "failed": int(reconciliation.get("failed", len(inputs.get("errors") or [])) or 0),
            "duplicate": int(reconciliation.get("duplicate", 0) or 0),
            "extracted": len(records),
        }
        expected = config.get("expectedScope") or config.get("expected_scope") or {}
        expected = expected if isinstance(expected, dict) else {}
        minimum = int(expected.get("minRecords", config.get("minimum_expected_records", 0)) or 0)
        if len(records) < minimum:
            errors.append({"code": "MINIMUM_EXPECTED_RECORDS", "expected": minimum, "actual": len(records)})
        if bool(expected.get("requireComplete", False)) and reconciliation["failed"]:
            errors.append({"code": "INCOMPLETE_SCOPE", "reconciliation": reconciliation})
        allowed_empty = bool(expected.get("allowEmpty", config.get("on_empty") == "allow"))
        if not records and not allowed_empty:
            errors.append({"code": "EMPTY_UNEXPECTED"})
        assessment_codes: list[str] = []
        if not records and allowed_empty and bool(inputs.get("source_checked", True)):
            assessment_codes.append("EMPTY_VALID_WINDOW")
        coverage = config.get("requiredFieldCoverage") or config.get("required_field_coverage") or {}
        if isinstance(coverage, dict) and records:
            for field, minimum_coverage in coverage.items():
                actual = sum(bool(isinstance(row, dict) and row.get(field) not in (None, "")) for row in records) / len(records)
                if actual < float(minimum_coverage):
                    errors.append({"code": "REQUIRED_FIELD_COVERAGE", "field": field, "expected": float(minimum_coverage), "actual": actual})
        source_role = config.get("sourceRole") or config.get("source_role") or {}
        if isinstance(source_role, dict) and source_role.get("expected"):
            actual_role = inputs.get("source_role") or inputs.get("sourceRole") or next(
                (
                    row.get("segment") or row.get("source_role")
                    for row in records
                    if isinstance(row, dict)
                    and (row.get("segment") or row.get("source_role"))
                ),
                None,
            )
            if actual_role != source_role["expected"]:
                errors.append({"code": "SOURCE_ROLE_MISMATCH", "expected": source_role["expected"], "actual": actual_role})
        expected_states = config.get("expectedStates") or config.get("expected_states") or {}
        if isinstance(expected_states, dict) and expected_states.get("states"):
            visited = {str(value) for value in (inputs.get("visited_states") or inputs.get("visitedStates") or [])}
            missing = [str(value) for value in expected_states["states"] if str(value) not in visited]
            if missing:
                errors.append({"code": "EXPECTED_STATE_MISSING", "expected": expected_states["states"], "visited": sorted(visited), "missing": missing})
        date_window = config.get("dateWindow") or config.get("date_window") or {}
        if isinstance(date_window, dict) and bool(date_window.get("forbidOutside", date_window.get("forbid_outside", False))):
            field = str(date_window.get("field") or "source_published_at")
            lower = _parse_effective_date(date_window.get("from"))
            upper = _parse_effective_date(date_window.get("to"))
            for index, row in enumerate(records):
                value = _parse_effective_date(row.get(field) if isinstance(row, dict) else None)
                if value is None or (lower and value < lower) or (upper and value >= upper):
                    errors.append({"code": "DATE_WINDOW_VIOLATION", "row": index, "field": field, "value": row.get(field) if isinstance(row, dict) else None})
        for key, code, denominator in (
            ("detailSuccessRatio", "DETAIL_COVERAGE_FAILED", reconciliation["discovered"]),
            ("documentParseRatio", "DOCUMENT_PARSE_INCOMPLETE", int(inputs.get("documents_discovered", 0) or 0)),
        ):
            assertion = config.get(key) or config.get(re.sub(r"(?<!^)(?=[A-Z])", "_", key).lower()) or {}
            if isinstance(assertion, dict) and assertion.get("min") is not None and denominator:
                succeeded = reconciliation["succeeded"] if key == "detailSuccessRatio" else int(inputs.get("documents_parsed", 0) or 0)
                actual = succeeded / denominator
                if actual < float(assertion["min"]):
                    errors.append({"code": code, "expected": float(assertion["min"]), "actual": actual})
        quarantine_enabled = bool(config.get("quarantine", config.get("quarantine_invalid", False)))
        if errors and quarantine_enabled:
            failed_rows = {int(error["row"]) for error in errors if isinstance(error.get("row"), int)}
            quarantined = [row for index, row in enumerate(records) if index in failed_rows]
            records = [row for index, row in enumerate(records) if index not in failed_rows]
        fail_required = str(config.get("errorPolicy") or config.get("error_policy") or "").upper() in {
            "FAIL", "FAIL_REQUIRED_SCOPE",
        }
        should_raise = bool(config.get("fail_on_error", True)) and int(config.get("contractVersion", 1)) != 2
        if errors and (should_raise or fail_required and int(config.get("contractVersion", 1)) != 2) and not quarantine_enabled:
            raise ValueError(f"Schema validation failed: {errors[:20]}")
        status = "PASS" if not errors else "PARTIAL" if records else "FAIL"
        return {
            "records": records,
            "valid": not errors,
            "errors": errors,
            "quarantined_records": quarantined,
            "count": len(records),
            "business_records": bool(inputs.get("business_records")),
            "assessment_status": status,
            "assessment_codes": assessment_codes + [str(error.get("code")) for error in errors if error.get("code")],
            "reconciliation": reconciliation,
            "commit_allowed": status == "PASS" or (status == "PARTIAL" and bool(config.get("allow_partial_commit", False))),
        }


class DeduplicateNode:
    type = "deduplicate"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        records = find_value(inputs, str(config.get("input_path", "records"))) or []
        keys = config.get("keys", [])
        seen: set[str] = set()
        output = []
        for row in records:
            key = json.dumps([row.get(k) for k in keys], sort_keys=True, ensure_ascii=False, default=str) if keys else json.dumps(row, sort_keys=True, ensure_ascii=False, default=str)
            if key not in seen:
                seen.add(key)
                output.append(row)
        return {"records": output, "count": len(output), "duplicates_removed": len(records) - len(output)}


class ConditionNode:
    type = "condition"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        left = find_value(inputs, str(config.get("field", "")))
        operation = str(config.get("operator", "eq"))
        right = config.get("value")
        operations = {
            "eq": lambda: left == right, "ne": lambda: left != right,
            "gt": lambda: left is not None and left > right, "gte": lambda: left is not None and left >= right,
            "lt": lambda: left is not None and left < right, "lte": lambda: left is not None and left <= right,
            "contains": lambda: right in left if left is not None else False,
            "exists": lambda: left is not None, "empty": lambda: left in (None, "", [], {}),
        }
        result = bool(operations.get(operation, lambda: False)())
        return {"condition": result, "true": inputs if result else None, "false": inputs if not result else None}


class OutputNode:
    type = "output"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        records = find_value(inputs, str(config.get("input_path", "records")))
        raw_records = records
        records = records if isinstance(records, list) else ([] if config.get("dataset_id") else raw_records)
        record_count = len(records) if isinstance(records, list) else (1 if records is not None else 0)
        explicit = bool(inputs.get("business_records"))
        minimum = max(0, int(config.get("minimum_expected_records", 0) or 0))
        on_empty = str(config.get("on_empty", "warning"))
        errors: list[dict[str, Any]] = list(inputs.get("mapping_errors") or [])
        assessment_status = str(inputs.get("assessment_status") or "PASS")
        commit_allowed = bool(inputs.get("commit_allowed", assessment_status == "PASS"))
        if not commit_allowed:
            errors.append({
                "code": "ASSESSMENT_BLOCKED",
                "assessment_status": assessment_status,
                "reconciliation": inputs.get("reconciliation", {}),
            })
        if not explicit and config.get("dataset_id"):
            errors.append({"code": "MAPPING_REQUIRED", "message": "Save Dataset принимает только business records из Mapping"})
        if record_count < minimum:
            errors.append({"code": "MINIMUM_EXPECTED_RECORDS", "expected": minimum, "actual": record_count})
        if not records and on_empty == "fail":
            raise ValueError("Save Dataset: пустой результат запрещён настройкой On empty = fail")
        return {"records": records, "count": record_count, "output_name": config.get("name", "result"),
                "business_records": explicit, "mapping_errors": inputs.get("mapping_errors", []),
                "preflight": {"input_records": record_count, "minimum_expected_records": minimum,
                              "on_empty": on_empty, "validation_errors": errors,
                              "assessment_status": assessment_status, "commit_allowed": not errors,
                              "schema_preview": records[:5] if isinstance(records, list) else records}}


class SaveExternalDatabaseNode:
    type = "save_external_db"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        from sqlalchemy import MetaData, Table, create_engine

        records = find_value(inputs, str(config.get("input_path", "records"))) or []
        records = records if isinstance(records, list) else [records]
        connection_name = str(config.get("connection") or "")
        connection = context.capabilities.get("database_connections", {}).get(connection_name)
        if not connection:
            raise ValueError(f"Подключение к БД не найдено: {connection_name}")
        table_name = str(config.get("table") or "")
        allowed_tables = connection.get("allowed_tables") or []
        qualified = f"{config.get('schema', 'public')}.{table_name}"
        if allowed_tables and table_name not in allowed_tables and qualified not in allowed_tables:
            raise ValueError(f"Таблица не входит в allowed_tables: {qualified}")
        if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", table_name):
            raise ValueError("Некорректное имя таблицы")
        schema_name = str(config.get("schema") or connection.get("schema") or "public")
        if connection["engine"] == "sqlite":
            schema_name = None
        engine = create_engine(connection["url"], pool_pre_ping=True)
        table = Table(table_name, MetaData(), schema=schema_name, autoload_with=engine)
        mapping = config.get("mapping") or {}
        prepared = [{mapping.get(key, key): value for key, value in row.items() if mapping.get(key, key) in table.c} for row in records if isinstance(row, dict)]
        if not prepared:
            return {"records": records, "written": 0}
        with engine.begin() as db_connection:
            mode = config.get("mode", "insert")
            if mode == "upsert" and config.get("conflict_keys"):
                if engine.dialect.name == "postgresql":
                    from sqlalchemy.dialects.postgresql import insert
                    statement = insert(table).values(prepared)
                    update_columns = {column.name: statement.excluded[column.name] for column in table.c if column.name not in config["conflict_keys"]}
                    db_connection.execute(statement.on_conflict_do_update(index_elements=config["conflict_keys"], set_=update_columns))
                elif engine.dialect.name in {"mysql", "mariadb"}:
                    from sqlalchemy.dialects.mysql import insert
                    statement = insert(table).values(prepared)
                    db_connection.execute(statement.on_duplicate_key_update(**{column.name: statement.inserted[column.name] for column in table.c if column.name not in config["conflict_keys"]}))
                else:
                    for row in prepared:
                        filters = [table.c[key] == row[key] for key in config["conflict_keys"]]
                        existing = db_connection.execute(table.select().where(*filters)).first()
                        if existing:
                            db_connection.execute(table.update().where(*filters).values(**row))
                        else:
                            db_connection.execute(table.insert().values(**row))
            else:
                db_connection.execute(table.insert(), prepared)
        return {"records": records, "written": len(prepared), "connection": connection_name, "table": qualified}


class ExportFileNode:
    type = "export_file"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        records = find_value(inputs, str(config.get("input_path", "records"))) or []
        records = records if isinstance(records, list) else [records]
        format_name = str(config.get("format", "xlsx")).lower()
        filename = str(config.get("filename") or f"export.{format_name}")
        if format_name == "json":
            data = json.dumps(records, ensure_ascii=False, indent=2, default=str).encode("utf-8")
            content_type = "application/json"
        elif format_name == "csv":
            output = io.StringIO()
            columns = sorted({key for row in records if isinstance(row, dict) for key in row})
            writer = csv.DictWriter(output, fieldnames=columns)
            writer.writeheader()
            writer.writerows(records)
            data = output.getvalue().encode("utf-8-sig")
            content_type = "text/csv"
        else:
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Данные"
            columns = sorted({key for row in records if isinstance(row, dict) for key in row})
            sheet.append(columns)
            for row in records:
                sheet.append([stringify_cell(row.get(column)) for column in columns])
            buffer = io.BytesIO()
            workbook.save(buffer)
            data = buffer.getvalue()
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        artifact = await store_artifact(context, data, content_type, "", filename, "export")
        return {"records": records, "export": artifact}


class SendWebhookNode:
    type = "send_webhook"

    async def execute(self, context: ExecutionContext, inputs: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
        url = render_template(str(config.get("url") or ""), context, inputs)
        if not url:
            raise ValueError("Webhook URL is required")
        payload = find_value(inputs, str(config.get("input_path", "records")))
        policy = FetchPolicy.from_config(config)
        async with httpx.AsyncClient(follow_redirects=False, timeout=policy.timeout) as client:
            response = await request_with_egress_policy(
                client, "POST", url, policy, egress_policy=EgressPolicy.from_config(config),
                json=payload, headers=render_object(config.get("headers") or {}, context, inputs),
            )
            response.raise_for_status()
        return {"sent": True, "status_code": response.status_code, "response": response.text[:2000], "records": payload, "redirect_chain": response.extensions.get("redirect_chain", [])}


NODE_REGISTRY = {node.type: node() for node in [
    ManualTriggerNode, HTTPRequestNode, BrowserOpenNode, DownloadFileNode, FollowLinksNode, PaginationNode,
    CrawlLinksNode,
    ParseHTMLNode, SelectElementsNode, ExtractRepeatingListNode, ParseTableNode, JSONPathNode, ParseDocumentNode,
    TransformNode, MappingNode, SetConstantNode, FormulaNode, LLMExtractNode, LLMClassifyNode, ValidateNode, DeduplicateNode, ConditionNode,
    OutputNode, SaveExternalDatabaseNode, ExportFileNode, SendWebhookNode,
]}


def find_value(data: Any, path: str) -> Any:
    if not path:
        return data
    current = data
    for part in re.findall(r"[^.\[\]]+", path):
        if isinstance(current, dict):
            current = current.get(part)
        elif isinstance(current, list) and part.isdigit():
            current = current[int(part)]
        else:
            return None
    return current


def simple_json_path(data: Any, path: str) -> list[Any]:
    """Small JSONPath fallback used when the optional parser is unavailable.

    Supports the picker output (object keys, numeric indices and ``[*]``) and
    deliberately returns no values for an absent path rather than ``[None]``.
    """
    if path in {"", "$"}:
        return [data]
    normalized = path.removeprefix("$").lstrip(".")
    tokens = re.findall(r"([^.[\]]+)|\[(\*|\d+)\]", normalized)
    if not tokens:
        return []
    values = [data]
    for key, index in tokens:
        next_values: list[Any] = []
        if key:
            for value in values:
                if isinstance(value, dict) and key in value:
                    next_values.append(value[key])
        elif index == "*":
            for value in values:
                if isinstance(value, list):
                    next_values.extend(value)
                elif isinstance(value, dict):
                    next_values.extend(value.values())
        else:
            for value in values:
                if isinstance(value, list) and int(index) < len(value):
                    next_values.append(value[int(index)])
        values = next_values
        if not values:
            break
    return values


def canonical_url(value: str, drop_query_params: list[str] | None = None) -> str:
    """Normalize a URL while preserving query parameters that carry identity."""
    parts = urlsplit(value)
    path = re.sub(r"/{2,}", "/", parts.path).rstrip("/")
    configured = {item.lower() for item in (drop_query_params or [])}
    tracking = {"utm", "gclid", "fbclid", "yclid", "mc_cid", "mc_eid"}
    query = [
        (key, item)
        for key, item in parse_qsl(parts.query, keep_blank_values=True)
        if key.lower() not in configured
        and key.lower() not in tracking
        and not key.lower().startswith("utm_")
    ]
    return urlunsplit((parts.scheme.lower(), parts.netloc.lower(), path, urlencode(sorted(query)), ""))


def build_url_frontier(
    items: list[Any],
    *,
    base_url: str,
    origin_url: str,
    url_path: str,
    config: dict[str, Any],
    limit: int,
) -> list[dict[str, Any]]:
    """Build one bounded canonical frontier for every link-following node.

    An empty ``allowed_domains`` list means that no cross-origin domains are
    added: the normal same-origin rule remains in force. Set
    ``same_origin_only=false`` explicitly to allow unrestricted domains.
    """
    pattern_text = str(config.get("url_pattern") or "").strip()
    pattern = re.compile(pattern_text, re.I) if pattern_text else None
    origin_host = (urlsplit(origin_url).hostname or "").lower()
    allowed_domains = {
        str(domain).strip().lower().split(":", 1)[0]
        for domain in (config.get("allowed_domains") or [])
        if str(domain).strip()
    }
    seen: set[str] = set()
    frontier: list[dict[str, Any]] = []
    for value in items:
        item = dict(value) if isinstance(value, dict) else {url_path: value}
        raw_url = find_value(item, url_path)
        if not raw_url:
            continue
        canonical = canonical_url(
            urljoin(base_url, str(raw_url)),
            list(config.get("drop_query_params") or []),
        )
        parts = urlsplit(canonical)
        if parts.scheme not in {"http", "https"}:
            continue
        candidate_host = (parts.hostname or "").lower()
        if allowed_domains and not any(
            candidate_host == domain or candidate_host.endswith(f".{domain}")
            for domain in allowed_domains
        ):
            continue
        if (
            not allowed_domains
            and config.get("same_origin_only", True)
            and origin_host
            and candidate_host != origin_host
        ):
            continue
        if pattern and not pattern.search(canonical):
            continue
        if canonical in seen:
            continue
        seen.add(canonical)
        frontier.append({"item": item, "url": canonical})
        if len(frontier) >= max(limit, 1):
            break
    return frontier


async def close_http_client(client: Any) -> None:
    close = getattr(client, "aclose", None)
    if close is not None:
        await close()


def crawl_resume_scope(context: ExecutionContext, config: dict[str, Any]) -> str:
    stable_config = {
        key: value
        for key, value in config.items()
        if key not in {"resume_token"} and not key.startswith("_force_")
    }
    fingerprint = hashlib.sha256(
        json.dumps(stable_config, sort_keys=True, separators=(",", ":"), default=str).encode()
    ).hexdigest()
    return f"{context.workflow_version_id}:{config.get('_node_id', 'crawl_links')}:{fingerprint}"


def crawl_resume_secret(context: ExecutionContext) -> bytes:
    secret = context.secrets.get("_CRAWL_RESUME_SECRET")
    if not secret:
        # Direct library use remains resumable; API execution injects the
        # deployment secret so production tokens cannot be forged.
        secret = f"multiverse-resume:{context.workflow_version_id}"
    return secret.encode()


def encode_crawl_resume_token(
    urls: list[str],
    context: ExecutionContext,
    config: dict[str, Any],
) -> str:
    payload = json.dumps(
        {"version": 1, "scope": crawl_resume_scope(context, config), "urls": urls},
        sort_keys=True,
        separators=(",", ":"),
    ).encode()
    signature = hmac.new(crawl_resume_secret(context), payload, hashlib.sha256).digest()
    body = base64.urlsafe_b64encode(payload).decode().rstrip("=")
    proof = base64.urlsafe_b64encode(signature).decode().rstrip("=")
    return f"{body}.{proof}"


def decode_crawl_resume_token(
    token: str,
    context: ExecutionContext,
    config: dict[str, Any],
) -> list[str]:
    try:
        body, proof = token.split(".", 1)
        payload = base64.urlsafe_b64decode(body + "=" * (-len(body) % 4))
        signature = base64.urlsafe_b64decode(proof + "=" * (-len(proof) % 4))
        expected = hmac.new(crawl_resume_secret(context), payload, hashlib.sha256).digest()
        if not hmac.compare_digest(signature, expected):
            raise ValueError
        value = json.loads(payload)
        urls = value.get("urls")
        if (
            value.get("version") != 1
            or value.get("scope") != crawl_resume_scope(context, config)
            or not isinstance(urls, list)
            or any(not isinstance(url, str) for url in urls)
        ):
            raise ValueError
        return [canonical_url(url) for url in urls]
    except Exception as exc:
        raise ValueError("Invalid crawl resume token") from exc


async def hydrate_dynamic_detail(
    client: httpx.AsyncClient, response: httpx.Response, config: dict[str, Any]
) -> tuple[str, bytes, str]:
    """Compatibility no-op for old imports.

    Dynamic rendering now happens through ``detail_fetch_mode`` in
    :class:`CrawlLinksNode`, which is portable across websites.
    """
    return response.text, response.content, response.headers.get("content-type", "text/html")


def extract_article_record(
    page_html: str,
    page_url: str,
    candidate: dict[str, Any],
    config: dict[str, Any],
    artifact: dict[str, Any] | None,
) -> dict[str, Any]:
    soup = BeautifulSoup(page_html, "lxml")
    detail_fields = config.get("detail_fields")
    if isinstance(detail_fields, list):
        record: dict[str, Any] = {"record_id": candidate.get("record_id") or hashlib.sha256(canonical_url(page_url).encode()).hexdigest()[:20]}
        page_metadata = extract_page_metadata(soup)
        if config.get("include_listing_fields") and isinstance(candidate.get("item"), dict):
            record.update(candidate["item"])
        for field in detail_fields:
            if not isinstance(field, dict) or not field.get("name"):
                continue
            name = str(field["name"])
            source_kind = str(field.get("source") or "selector").lower()
            if source_kind in {"listing", "item"}:
                listing_item = candidate.get("item") if isinstance(candidate.get("item"), dict) else {}
                value = find_value(listing_item, str(field.get("source_path") or name))
                if name in {"source_published_at", "source_modified_at"}:
                    value = normalize_source_datetime(
                        str(value or ""),
                        timezone=str(field.get("timezone") or config.get("timezone") or "UTC"),
                        date_format=str(field.get("format") or ""),
                    )
                record[name] = value
                continue
            if source_kind in {"metadata", "json_ld"}:
                metadata_key = str(field.get("metadata_key") or name)
                value = page_metadata.get(metadata_key)
                if name in {"source_published_at", "source_modified_at"}:
                    value = normalize_source_datetime(
                        str(value or ""),
                        timezone=str(field.get("timezone") or config.get("timezone") or "UTC"),
                        date_format=str(field.get("format") or ""),
                    )
                record[name] = value
                continue
            if source_kind in {"response", "detail_response", "json"}:
                response_payload = candidate.get("detail_response") if isinstance(candidate.get("detail_response"), dict) else {}
                value = find_value(response_payload, str(field.get("source_path") or name))
                value_mode = str(field.get("value") or "raw").lower()
                if value_mode == "html_text":
                    value = BeautifulSoup(str(value or ""), "lxml").get_text(" ", strip=True)
                elif value_mode == "join" and isinstance(value, list):
                    value = str(field.get("separator") or "|").join(clean_inline_text(str(item)) for item in value if item is not None)
                elif value_mode == "json":
                    value = json.dumps(value, ensure_ascii=False)
                regex = str(field.get("regex") or "")
                if regex and value is not None:
                    match = re.search(regex, str(value))
                    group = int(field.get("regex_group", 1))
                    value = match.group(group) if match else None
                if name in {"source_published_at", "source_modified_at"}:
                    value = normalize_source_datetime(
                        str(value or ""),
                        timezone=str(field.get("timezone") or config.get("timezone") or "UTC"),
                        date_format=str(field.get("format") or ""),
                    )
                record[name] = value
                continue
            selector = str(field.get("selector") or "")
            elements = soup.select(selector) if selector else []
            element = elements[0] if elements else None
            value_mode = str(field.get("value") or "text")
            if field.get("multiple") and value_mode == "links":
                links = [
                    {
                        "title": item.get_text(" ", strip=True) or Path(urlsplit(str(item.get("href") or "")).path).name,
                        "url": canonical_url(urljoin(page_url, str(item.get("href") or ""))),
                    }
                    for item in elements if item.get("href")
                ]
                # The legacy ``attachments_json`` field intentionally remains
                # a JSON string for backwards compatibility.  The shared
                # market-news schema uses the renamed ``attachments`` field,
                # which is declared as an array and must stay structured all
                # the way through Mapping → persistence validation.
                record[name] = links if name == "attachments" else json.dumps(links, ensure_ascii=False)
                continue
            if element is None:
                record[name] = None
                continue
            attribute = str(field.get("attribute") or "")
            value = element.get(attribute) if attribute else (
                element.decode_contents() if value_mode == "html" else element.get_text(" ", strip=True)
            )
            if name in {"source_published_at", "source_modified_at"}:
                value = normalize_source_datetime(
                    str(value or ""),
                    timezone=str(field.get("timezone") or config.get("timezone") or "UTC"),
                    date_format=str(field.get("format") or ""),
                )
            record[name] = value
        constants = config.get("detail_constants") or {}
        if isinstance(constants, dict):
            record.update(constants)
        record["fetched_at"] = candidate.get("fetched_at") or datetime.now(UTC).isoformat()
        record["url"] = canonical_url(page_url, list(config.get("drop_query_params") or []))
        if artifact:
            record["__provenance"] = {"raw_artifact": artifact}
        return record
    title_selector = str(config.get("title_selector") or "")
    title = select_text(soup, title_selector) if title_selector else ""
    title = title or str(candidate["item"].get("title") or "")
    if not title:
        title = next((select_text(soup, selector) for selector in ("h1", "[itemprop='headline']", "article h1", "meta[property='og:title']") if select_text(soup, selector)), "")
    listing_date = str(candidate["item"].get("shortDate") or candidate["item"].get("published_at") or candidate["item"].get("date") or "")
    date_selector = str(config.get("date_selector") or "")
    date_text = listing_date if re.search(r"\d{4}-\d{2}-\d{2}", listing_date) else select_text(soup, date_selector) if date_selector else ""
    if not date_text:
        date_text = next((element.get("datetime") or element.get("content") or element.get_text(" ", strip=True) for element in soup.select("time, [itemprop='datePublished'], meta[property='article:published_time'], meta[name='date']") if element.get("datetime") or element.get("content") or element.get_text(" ", strip=True)), "")
    jsonld_dates = extract_jsonld_dates(soup)
    if not date_text:
        date_text = jsonld_dates.get("source_published_at", "")
    body_selector = str(config.get("body_selector") or "")
    body = soup.select_one(body_selector) if body_selector else None
    if not body:
        body = next((soup.select_one(selector) for selector in ("[itemprop='articleBody']", "article", "main article", ".article-body", ".post-content", ".entry-content", "main") if soup.select_one(selector)), None)
    body_html = body.decode_contents() if body else ""
    body_text = clean_article_text(body.get_text("\n", strip=True) if body else "")
    # A detail page may legitimately omit the preferred article wrapper (for
    # example a short technical notice). Keep the universal contract stable:
    # every discovered card still receives auditable text and HTML evidence.
    if not body_text:
        fallback = soup.select_one("main") or soup.body
        if fallback:
            body_html = body_html or fallback.decode_contents()
            body_text = clean_article_text(fallback.get_text("\n", strip=True))
    if not body_text:
        body_text = clean_article_text(title) or clean_article_text(page_url)
    if not body_html:
        body_html = f"<p>{escape(body_text)}</p>"
    tag_selector = str(config.get("tag_selector") or "")
    tags = unique_strings(element.get_text(" ", strip=True) for element in soup.select(tag_selector)) if tag_selector else unique_strings([element.get("content") or element.get_text(" ", strip=True) for element in soup.select("[data-parser-studio-tag], [rel='tag'], meta[name='keywords']")])
    attachment_selector = str(config.get("attachment_selector") or "a[href$='.pdf'],a[href$='.doc'],a[href$='.docx'],a[href$='.xls'],a[href$='.xlsx'],a[href$='.zip']")
    attachments = []
    if body:
        for element in body.select(attachment_selector):
            href = element.get("href")
            if href:
                attachments.append({"title": element.get_text(" ", strip=True) or Path(urlsplit(href).path).name, "url": canonical_url(urljoin(page_url, href))})
    published_at = (
        normalize_source_datetime(date_text)
        if re.search(r"T\d{2}:\d{2}(?::\d{2})?(?:Z|[+-]\d{2}:?\d{2})$", date_text)
        else normalize_publication_date(date_text)
    )
    record: dict[str, Any] = {
        "record_id": candidate.get("record_id") or candidate.get("news_id") or hashlib.sha256(canonical_url(page_url).encode("utf-8")).hexdigest()[:20],
        "title": clean_inline_text(title),
        "published_at": published_at or "",
        "url": canonical_url(page_url),
        "body_text": body_text,
        "body_html": body_html,
        "tags": "|".join(tags),
        "attachments_json": json.dumps(attachments, ensure_ascii=False),
        "language": str(config.get("language") or (soup.html or {}).get("lang") or "").split("-", 1)[0],
        "source_name": str(config.get("source_name") or candidate.get("item", {}).get("source_name") or ""),
        "fetched_at": candidate.get("fetched_at") or datetime.now(UTC).isoformat(),
        "observed_at": datetime.now(UTC).isoformat(),
    }
    if jsonld_dates.get("source_modified_at"):
        record["source_modified_at"] = normalize_source_datetime(
            jsonld_dates["source_modified_at"]
        )
    # Older saved workflows may have supplied an additional domain identifier.
    # Preserve it as data, without making the crawler or the system template
    # depend on a particular field name.
    if candidate.get("news_id"):
        record["news_id"] = candidate["news_id"]
    if artifact:
        record["__provenance"] = {"raw_artifact": artifact}
    return record


def extract_jsonld_dates(soup: BeautifulSoup) -> dict[str, str]:
    """Read generic Schema.org dates without assuming a site or industry."""
    result: dict[str, str] = {}

    def visit(value: Any) -> None:
        if isinstance(value, dict):
            if not result.get("source_published_at") and value.get("datePublished"):
                result["source_published_at"] = str(value["datePublished"])
            if not result.get("source_modified_at") and value.get("dateModified"):
                result["source_modified_at"] = str(value["dateModified"])
            for child in value.values():
                if isinstance(child, (dict, list)):
                    visit(child)
        elif isinstance(value, list):
            for child in value:
                visit(child)

    for script in soup.select('script[type="application/ld+json"]'):
        try:
            visit(json.loads(script.string or script.get_text() or ""))
        except (json.JSONDecodeError, TypeError):
            continue
    return result


def extract_page_metadata(soup: BeautifulSoup) -> dict[str, str]:
    result = extract_jsonld_dates(soup)
    semantic = {
        "source_published_at": (
            "time[datetime]",
            "[itemprop='datePublished']",
            "meta[property='article:published_time']",
        ),
        "source_modified_at": (
            "[itemprop='dateModified']",
            "meta[property='article:modified_time']",
        ),
    }
    for target, selectors in semantic.items():
        if result.get(target):
            continue
        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                value = element.get("datetime") or element.get("content") or element.get_text(" ", strip=True)
                if value:
                    result[target] = str(value)
                    break
    return result


def normalize_source_datetime(value: str, *, timezone: str = "UTC", date_format: str = "") -> str | None:
    text = clean_inline_text(value)
    if not text:
        return None
    translated_format = date_format
    for token, replacement in (
        ("YYYY", "%Y"), ("DD", "%d"), ("MM", "%m"),
        ("HH", "%H"), ("mm", "%M"), ("ss", "%S"),
    ):
        translated_format = translated_format.replace(token, replacement)
    parsed: datetime | None = None
    if translated_format:
        try:
            parsed = datetime.strptime(text, translated_format)
        except ValueError:
            return None
    else:
        try:
            parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            normalized = normalize_publication_date(text)
            try:
                parsed = datetime.fromisoformat(normalized) if normalized != text else None
            except ValueError:
                parsed = None
    if parsed is None:
        return None
    if parsed.tzinfo is None:
        try:
            parsed = parsed.replace(tzinfo=ZoneInfo(timezone))
        except Exception as exc:
            raise ValueError(f"Unknown source timezone: {timezone}") from exc
    try:
        return parsed.astimezone(UTC).isoformat().replace("+00:00", "Z")
    except (OverflowError, ValueError):
        return None


def select_text(soup: BeautifulSoup, selector: str) -> str:
    return soup.select_one(selector).get_text(" ", strip=True) if selector and soup.select_one(selector) else ""


def clean_inline_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def clean_article_text(value: str) -> str:
    lines = [clean_inline_text(line) for line in value.replace("\r", "").split("\n")]
    return "\n\n".join(line for line in lines if line)


def unique_strings(values: Any) -> list[str]:
    result: list[str] = []
    for value in values:
        text = clean_inline_text(str(value))
        if text and text not in result:
            result.append(text)
    return result


def normalize_publication_date(value: str) -> str:
    value = clean_inline_text(value)
    if not value:
        return ""
    iso = re.search(r"\d{4}-\d{2}-\d{2}(?:T\d{2}:\d{2}:\d{2})?", value)
    if iso:
        return iso.group(0)
    months = {"января": 1, "февраля": 2, "марта": 3, "апреля": 4, "мая": 5, "июня": 6, "июля": 7, "августа": 8, "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12}
    match = re.search(r"(\d{1,2})\s+([а-яё]+)\s+(\d{4})", value.lower())
    if match and match.group(2) in months:
        return f"{match.group(3)}-{months[match.group(2)]:02d}-{int(match.group(1)):02d}"
    return value


def render_template(template: str, context: ExecutionContext, inputs: dict[str, Any]) -> str:
    values: dict[str, Any] = {"run.id": context.run_id, **context.variables, **inputs}

    def replacement(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        if key.startswith("secret."):
            return context.secrets.get(key[7:], "")
        value = find_value(values, key)
        if value is None and key.startswith("input."):
            value = find_value(inputs, key[6:])
        return "" if value is None else str(value)

    return re.sub(r"{{\s*([^}]+)\s*}}", replacement, template)


def render_object(value: Any, context: ExecutionContext, inputs: dict[str, Any]) -> Any:
    if isinstance(value, str):
        return render_template(value, context, inputs)
    if isinstance(value, list):
        return [render_object(item, context, inputs) for item in value]
    if isinstance(value, dict):
        return {key: render_object(item, context, inputs) for key, item in value.items()}
    return value


async def response_payload(context: ExecutionContext, response: httpx.Response) -> dict[str, Any]:
    content_type = response.headers.get("content-type", "")
    raw = response.content
    artifact = await store_artifact(context, raw, content_type or "application/octet-stream", str(response.url), filename_from_response(response), "raw_document")
    if "json" in content_type:
        payload: Any = response.json()
    elif any(token in content_type for token in ("text", "html", "xml", "javascript")) or not content_type:
        # A response can declare UTF-8 while serving a legacy Cyrillic
        # encoding. Prefer the decoding with more readable text.
        payload = response.text
        try:
            cp1251_payload = raw.decode("cp1251")
            def readable(text: str) -> int:
                return sum("А" <= char <= "я" or char in "ЁёЎўІі" for char in text)
            if readable(cp1251_payload) > readable(payload) * 2:
                payload = cp1251_payload
        except UnicodeDecodeError:
            pass
    else:
        payload = base64.b64encode(raw).decode("ascii")
    arrays = json_array_paths(payload) if "json" in content_type else []
    return {"url": str(response.url), "status_code": response.status_code, "headers": dict(response.headers), "content_type": content_type, "body": payload, "content_base64": base64.b64encode(raw).decode("ascii"), "filename": filename_from_response(response), "sha256": artifact["sha256"], "artifact": artifact,
            "document_diagnostics": {"status": response.status_code, "content_type": content_type, "body_size": len(raw),
                                     "document_preview": stringify(payload)[:1000], "json_detected": bool(arrays) or "json" in content_type,
                                     "array_paths": arrays}}


def json_array_paths(value: Any, path: str = "$") -> list[dict[str, Any]]:
    paths: list[dict[str, Any]] = []
    if isinstance(value, list):
        paths.append({"path": f"{path}[*]", "length": len(value)})
        if value:
            paths.extend(json_array_paths(value[0], f"{path}[*]"))
    elif isinstance(value, dict):
        for key, child in value.items():
            paths.extend(json_array_paths(child, f"{path}.{key}"))
    return paths[:30]


async def store_artifact(context: ExecutionContext, data: bytes, content_type: str, url: str, filename: str, kind: str) -> dict[str, Any]:
    secret_values = list(context.secrets.values())
    data = redact_artifact_bytes(data, content_type, secret_values)
    sha256 = hashlib.sha256(data).hexdigest()
    safe_filename = re.sub(r"[^a-zA-Z0-9._-]+", "_", redact_text(filename or "artifact.bin", secret_values))
    artifact: dict[str, Any] = {"kind": kind, "url": redact_text(url, secret_values), "sha256": sha256, "content_type": content_type, "size": len(data), "filename": safe_filename}
    if context.artifact_storage is not None:
        stored = await context.artifact_storage.put_bytes("raw" if kind != "export" else "exports", f"runs/{context.run_id}/{sha256}-{safe_filename}", data, content_type, {"run_id": context.run_id})
        artifact.update(stored)
    context.artifacts.append(artifact)
    return artifact


def filename_from_response(response: httpx.Response) -> str:
    disposition = response.headers.get("content-disposition", "")
    match = re.search(r"filename\*?=(?:UTF-8''|\")?([^\";]+)", disposition, re.I)
    if match:
        return Path(match.group(1)).name
    name = Path(response.url.path).name
    return name or "document.bin"


def apply_operation(row: dict[str, Any], operation: dict[str, Any], *, context: ExecutionContext | None = None) -> None:
    field = operation.get("field")
    kind = operation.get("type")
    if kind == "rename":
        row[operation["to"]] = row.pop(field, None)
    elif kind == "constant":
        row[field] = operation.get("value")
    elif kind == "add_context":
        _apply_context_fields(row, operation, context)
    elif kind == "copy":
        target = str(operation.get("to") or field or "")
        if target:
            row[target] = find_value(row, str(operation.get("source") or operation.get("from") or ""))
    elif kind == "coalesce":
        target = str(operation.get("to") or field or "")
        values = operation.get("fields") or operation.get("sources") or []
        if target and isinstance(values, list):
            row[target] = next((find_value(row, str(source)) for source in values if find_value(row, str(source)) not in (None, "")), operation.get("default"))
    elif kind == "trim" and row.get(field) is not None:
        row[field] = str(row[field]).strip()
    elif kind == "normalize_spaces" and row.get(field) is not None:
        row[field] = re.sub(r"\s+", " ", str(row[field])).strip()
    elif kind == "replace" and row.get(field) is not None:
        row[field] = str(row[field]).replace(str(operation.get("search", "")), str(operation.get("replacement", "")))
    elif kind == "regex" and row.get(field) is not None:
        match = re.search(str(operation["pattern"]), str(row[field]), flags=re.I if "i" in str(operation.get("flags", "")) else 0)
        row[field] = match.group(int(operation.get("group", 0))) if match else operation.get("default")
    elif kind == "number":
        value = normalize_number(row.get(field))
        row[field] = float(value) if value is not None else None
    elif kind == "integer":
        value = normalize_number(row.get(field))
        row[field] = int(value) if value is not None else None
    elif kind == "currency":
        row[field] = normalize_currency(str(row.get(field, "")))
    elif kind == "term":
        row.update(normalize_term(str(row.get(field, ""))))
    elif kind == "rate":
        row.update(parse_rate_expression(str(row.get(field, ""))))
    elif kind == "map":
        row[field] = operation.get("mapping", {}).get(str(row.get(field)), row.get(field))
    elif kind == "split" and row.get(field) is not None:
        row[operation.get("to", field)] = str(row[field]).split(str(operation.get("separator", ",")))
    elif kind == "concat":
        row[field] = str(operation.get("separator", " ")).join(str(row.get(source, "")) for source in operation.get("fields", []))
    elif kind == "select_by_rules":
        _apply_selection_rules(row, operation)
    elif kind == "classify_access":
        _apply_access_rules(row, operation)


def _apply_context_fields(row: dict[str, Any], operation: Mapping[str, Any], context: ExecutionContext | None) -> None:
    """Inject run-context values into every record (universal ``add_context``).

    Values come from the execution context (source binding, run clock) and the
    record's own transport provenance.  Existing non-empty values always win,
    so an explicit mapping is never silently overwritten.
    """

    fields = operation.get("fields") or ["source_id", "source_name", "fetched_at", "page_url", "state"]
    if isinstance(fields, str):
        fields = [item.strip() for item in fields.split(",") if item.strip()]
    variables = context.variables if context is not None and isinstance(context.variables, Mapping) else {}
    source = variables.get("source") if isinstance(variables.get("source"), Mapping) else {}
    provenance = row.get("__provenance") if isinstance(row.get("__provenance"), Mapping) else {}
    page_provenance = provenance.get("page") if isinstance(provenance.get("page"), Mapping) else {}
    for name in fields:
        name = str(name)
        if name in row and row[name] not in (None, ""):
            continue
        if name == "source_id":
            value = source.get("id")
        elif name == "source_name":
            value = source.get("name")
        elif name == "fetched_at":
            value = (context.effective_run_clock if context is not None else None) or datetime.now(UTC)
        elif name == "page_url":
            value = row.get("url") or row.get("page_url")
        elif name == "state":
            value = row.get("state") or page_provenance.get("state") or provenance.get("state")
        else:
            value = source.get(name) or variables.get(name)
        if isinstance(value, datetime):
            value = value.isoformat()
        if value not in (None, ""):
            row[name] = value


def _rule_text(row: Mapping[str, Any], fields: Any) -> str:
    selected = fields if isinstance(fields, list) else [fields] if fields else list(row)
    return "\n".join(str(find_value(dict(row), str(field)) or "") for field in selected)


def _matches_rule(row: Mapping[str, Any], when: Any, default_fields: Any) -> tuple[bool, list[str]]:
    """Evaluate transparent text rules stored in a preset revision.

    Rule evaluation intentionally stays small and portable: case-insensitive
    literal or regular-expression patterns over operator-selected fields.  It
    is suitable for visible include/exclude/access policies but is never a
    hidden classifier or a source-specific branch in the engine.
    """

    spec = dict(when) if isinstance(when, Mapping) else {}
    text = _rule_text(row, spec.get("fields", default_fields))
    flags = re.I if bool(spec.get("ignoreCase", spec.get("ignore_case", True))) else 0

    def patterns(key: str) -> list[str]:
        value = spec.get(key, [])
        return [str(item) for item in (value if isinstance(value, list) else [value]) if str(item)]

    def matched(pattern: str) -> bool:
        try:
            return bool(re.search(pattern, text, flags)) if bool(spec.get("regex", False)) else pattern.casefold() in text.casefold()
        except re.error:
            return False

    any_patterns = patterns("anyPatterns") or patterns("any_patterns")
    all_patterns = patterns("allPatterns") or patterns("all_patterns")
    none_patterns = patterns("nonePatterns") or patterns("none_patterns")
    matched_terms = [pattern for pattern in any_patterns + all_patterns + none_patterns if matched(pattern)]
    return (
        (not any_patterns or any(matched(pattern) for pattern in any_patterns))
        and all(matched(pattern) for pattern in all_patterns)
        and not any(matched(pattern) for pattern in none_patterns),
        matched_terms,
    )


def _apply_selection_rules(row: dict[str, Any], operation: Mapping[str, Any]) -> None:
    default = operation.get("default") if isinstance(operation.get("default"), Mapping) else {}
    decision = {
        "action": str(default.get("action") or "AMBIGUOUS").upper(),
        "id": str(default.get("ruleId") or default.get("rule_id") or "unclassified-v1"),
        "reason": str(default.get("reason") or "no deterministic rule matched"),
        "matched": [],
    }
    for raw_rule in operation.get("rules") or []:
        if not isinstance(raw_rule, Mapping):
            continue
        applies, matched = _matches_rule(row, raw_rule.get("when"), operation.get("fields"))
        if applies:
            decision = {
                "action": str(raw_rule.get("action") or "AMBIGUOUS").upper(),
                "id": str(raw_rule.get("id") or "unnamed-rule"),
                "reason": str(raw_rule.get("reason") or raw_rule.get("id") or "configured selection rule"),
                "matched": matched,
            }
            break
    row["candidate_status"] = decision["action"] if decision["action"] in {"INCLUDE", "EXCLUDE", "AMBIGUOUS"} else "AMBIGUOUS"
    row["selection_rule_id"] = decision["id"]
    row["selection_reason"] = decision["reason"]
    row["selection_evidence"] = {"matched_terms": decision["matched"], "fields": operation.get("fields") or []}


def _apply_access_rules(row: dict[str, Any], operation: Mapping[str, Any]) -> None:
    status = str(operation.get("default") or "PUBLIC").upper()
    matched_terms: list[str] = []
    for raw_rule in operation.get("rules") or []:
        if not isinstance(raw_rule, Mapping):
            continue
        applies, matched = _matches_rule(row, raw_rule.get("when"), operation.get("fields"))
        if applies:
            status = str(raw_rule.get("status") or raw_rule.get("action") or status).upper()
            matched_terms = matched
            break
    row["access_status"] = status if status in {"PUBLIC", "PAYWALLED", "ACCESS_LIMITED"} else "ACCESS_LIMITED"
    row["access_evidence"] = {"matched_terms": matched_terms, "fields": operation.get("fields") or []}
    if row["access_status"] != "PUBLIC":
        for field in operation.get("redactFields") or operation.get("redact_fields") or []:
            row[str(field)] = None


_COLLECTION_OPERATION_TYPES = frozenset({"explode", "matrix_to_records", "unpivot", "expand_tiers", "select_effective_revision"})


def apply_collection_operation(
    records: list[dict[str, Any]], operation: dict[str, Any], context: ExecutionContext
) -> list[dict[str, Any]]:
    """Apply a declarative records-to-records transformation.

    The helper deliberately knows only structural configuration.  It has no
    source names, URLs or markup assumptions, so presets can safely use it for
    tables, API payloads and parsed documents alike.
    """

    kind = str(operation.get("type") or "")
    if kind == "explode":
        source = str(operation.get("field") or operation.get("source") or "")
        target = str(operation.get("target") or operation.get("to") or source)
        result: list[dict[str, Any]] = []
        for index, row in enumerate(records):
            values = row.get(source)
            values = values if isinstance(values, list) else ([] if values is None else [values])
            for value_index, value in enumerate(values):
                item = dict(row)
                item[target] = value
                _collection_provenance(item, kind, index, source, value_index)
                result.append(item)
        return result
    if kind in {"matrix_to_records", "unpivot"}:
        dimensions = operation.get("dimensionColumns") or operation.get("dimension_columns") or {}
        selector = str(dimensions.get("selector") or "*")
        header_target = str(dimensions.get("headerTarget") or dimensions.get("header_target") or "dimension_raw")
        value_target = str(dimensions.get("valueTarget") or dimensions.get("value_target") or "value_raw")
        id_fields = operation.get("idFields") or operation.get("id_fields") or []
        skip_empty = bool(operation.get("skipEmpty", operation.get("skip_empty", False)))
        result = []
        for index, row in enumerate(records):
            base = {field: row.get(field) for field in id_fields}
            for column, value in row.items():
                if not fnmatchcase(str(column), selector):
                    continue
                if skip_empty and value in (None, "", [], {}):
                    continue
                item = {**base, header_target: str(column).split(":", 1)[-1], value_target: value}
                if isinstance(row.get("evidence"), dict):
                    item["evidence"] = dict(row["evidence"])
                _collection_provenance(item, kind, index, str(column))
                result.append(item)
        return result
    if kind == "expand_tiers":
        source = str(operation.get("field") or operation.get("source") or "tiers")
        result = []
        for index, row in enumerate(records):
            tiers = row.get(source)
            tiers = tiers if isinstance(tiers, list) else ([] if tiers is None else [tiers])
            for tier_index, tier in enumerate(tiers):
                item = {key: value for key, value in row.items() if key != source}
                if isinstance(tier, dict):
                    item.update(tier)
                else:
                    item[str(operation.get("target") or source)] = tier
                _collection_provenance(item, kind, index, source, tier_index)
                result.append(item)
        return result
    if kind == "select_effective_revision":
        at = context.effective_run_clock or datetime.now(UTC)
        from_field = str(operation.get("effectiveFromField") or operation.get("effective_from_field") or "effective_from")
        to_field = str(operation.get("effectiveToField") or operation.get("effective_to_field") or "effective_to")
        candidates = [row for row in records if _effective_candidate(row, from_field, to_field, at)]
        if not candidates:
            return records
        selected = max(candidates, key=lambda item: _parse_effective_date(item.get(from_field)) or datetime.min.replace(tzinfo=UTC))
        selected = dict(selected)
        selected.setdefault("__provenance", {})["effective_revision"] = {
            "operation": kind,
            "effective_at": at.isoformat(),
            "candidate_count": len(records),
            "decision_rule": "current_then_latest_non_future",
        }
        return [selected]
    return records


def _collection_provenance(item: dict[str, Any], operation: str, source_row: int, source_column: str, item_index: int | None = None) -> None:
    provenance = item.get("__provenance") if isinstance(item.get("__provenance"), dict) else {}
    collection = {"operation": operation, "source_row": source_row, "source_column": source_column}
    if item_index is not None:
        collection["source_item"] = item_index
    item["__provenance"] = {**provenance, "collection": collection}


def _parse_effective_date(value: Any) -> datetime | None:
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return None
    return parsed.replace(tzinfo=UTC) if parsed.tzinfo is None else parsed.astimezone(UTC)


def _effective_candidate(row: dict[str, Any], from_field: str, to_field: str, at: datetime) -> bool:
    start, end = _parse_effective_date(row.get(from_field)), _parse_effective_date(row.get(to_field))
    return (start is None or start <= at) and (end is None or at < end)


def normalize_table_field_name(header: Any) -> str:
    """Create a stable structural key without guessing business meaning."""
    text = re.sub(r"[^\w]+", "_", str(header or "").strip().casefold(), flags=re.UNICODE)
    return text.strip("_")


def safe_eval(expression: str, values: dict[str, Any], run_clock: datetime | None = None) -> Any:
    tree = ast.parse(expression, mode="eval")

    def calculate(node: ast.AST) -> Any:
        if isinstance(node, ast.Expression):
            return calculate(node.body)
        if isinstance(node, ast.Constant):
            return node.value
        if isinstance(node, ast.Name):
            return values.get(node.id)
        if isinstance(node, ast.Call):
            if not isinstance(node.func, ast.Name):
                raise ValueError("Разрешены только встроенные функции Formula")
            function = FORMULA_FUNCTIONS.get(node.func.id)
            if function is None:
                raise ValueError(f"Неизвестная функция Formula: {node.func.id}")
            if node.keywords:
                raise ValueError("Именованные аргументы Formula не поддерживаются")
            arguments = [calculate(argument) for argument in node.args]
            if node.func.id in {"now", "today", "yesterday"}:
                return function(*arguments, run_clock=run_clock)
            return function(*arguments)
        if isinstance(node, ast.UnaryOp):
            value = calculate(node.operand)
            if isinstance(node.op, ast.USub):
                return -value
            if isinstance(node.op, ast.UAdd):
                return +value
            raise ValueError("Недопустимый unary operator")
        if isinstance(node, ast.BinOp):
            left, right = calculate(node.left), calculate(node.right)
            operations = {ast.Add: lambda: left + right, ast.Sub: lambda: left - right, ast.Mult: lambda: left * right, ast.Div: lambda: left / right, ast.Mod: lambda: left % right, ast.Pow: lambda: left ** right}
            operation = operations.get(type(node.op))
            if operation is None:
                raise ValueError("Недопустимый binary operator")
            return operation()
        if isinstance(node, ast.BoolOp):
            values_result = [bool(calculate(item)) for item in node.values]
            return all(values_result) if isinstance(node.op, ast.And) else any(values_result)
        if isinstance(node, ast.Compare):
            left = calculate(node.left)
            for operator, comparator in zip(node.ops, node.comparators, strict=True):
                right = calculate(comparator)
                comparisons = {ast.Eq: left == right, ast.NotEq: left != right, ast.Gt: left > right, ast.GtE: left >= right, ast.Lt: left < right, ast.LtE: left <= right}
                if type(operator) not in comparisons or not comparisons[type(operator)]:
                    return False
                left = right
            return True
        if isinstance(node, ast.IfExp):
            return calculate(node.body) if calculate(node.test) else calculate(node.orelse)
        raise ValueError(f"Expression содержит запрещённую операцию: {type(node).__name__}")

    return calculate(tree)


def formula_timezone(value: str) -> ZoneInfo:
    try:
        return ZoneInfo(str(value))
    except Exception as exc:
        raise ValueError(f"Некорректный IANA timezone: {value}") from exc


def formula_datetime(value: Any, timezone_name: str | None = None) -> datetime:
    if isinstance(value, datetime):
        parsed = value
    elif isinstance(value, date):
        parsed = datetime.combine(value, datetime_time.min)
    elif isinstance(value, str):
        normalized = value.replace("Z", "+00:00")
        try:
            parsed = datetime.fromisoformat(normalized)
        except ValueError:
            parsed = datetime.combine(date.fromisoformat(value), datetime_time.min)
    else:
        raise ValueError("Дата должна быть ISO-строкой или датой")
    if timezone_name:
        zone = formula_timezone(timezone_name)
        return parsed.replace(tzinfo=zone) if parsed.tzinfo is None else parsed.astimezone(zone)
    return parsed


def formula_now(timezone_name: str, run_clock: datetime | None = None) -> str:
    current = run_clock or datetime.now(formula_timezone(timezone_name))
    return current.astimezone(formula_timezone(timezone_name)).isoformat()


def formula_today(timezone_name: str, run_clock: datetime | None = None) -> str:
    current = run_clock or datetime.now(formula_timezone(timezone_name))
    return current.astimezone(formula_timezone(timezone_name)).date().isoformat()


def formula_yesterday(timezone_name: str, run_clock: datetime | None = None) -> str:
    current = run_clock or datetime.now(formula_timezone(timezone_name))
    return (current.astimezone(formula_timezone(timezone_name)).date() - timedelta(days=1)).isoformat()


def formula_start_of_day(value: Any, timezone_name: str) -> str:
    parsed = formula_datetime(value, timezone_name)
    return datetime.combine(parsed.date(), datetime_time.min, formula_timezone(timezone_name)).isoformat()


def formula_end_of_day(value: Any, timezone_name: str) -> str:
    parsed = formula_datetime(value, timezone_name)
    return datetime.combine(parsed.date(), datetime_time.max, formula_timezone(timezone_name)).isoformat()


def formula_add_days(value: Any, days: Any) -> str:
    parsed = formula_datetime(value)
    result = parsed + timedelta(days=int(days))
    return result.date().isoformat() if isinstance(value, str) and "T" not in value else result.isoformat()


def formula_format_date(value: Any, pattern: str) -> str:
    parsed = formula_datetime(value)
    replacements = {"YYYY": "%Y", "MM": "%m", "DD": "%d", "HH": "%H", "mm": "%M", "ss": "%S"}
    python_pattern = str(pattern)
    for token, directive in replacements.items(): python_pattern = python_pattern.replace(token, directive)
    return parsed.strftime(python_pattern)


def formula_parse_date(value: Any, pattern: str) -> str:
    replacements = {"YYYY": "%Y", "MM": "%m", "DD": "%d", "HH": "%H", "mm": "%M", "ss": "%S"}
    python_pattern = str(pattern)
    for token, directive in replacements.items(): python_pattern = python_pattern.replace(token, directive)
    parsed = datetime.strptime(str(value), python_pattern)
    return parsed.date().isoformat() if all(token not in str(pattern) for token in ("HH", "mm", "ss")) else parsed.isoformat()


FORMULA_FUNCTIONS = {
    "now": formula_now, "today": formula_today, "yesterday": formula_yesterday,
    "start_of_day": formula_start_of_day, "end_of_day": formula_end_of_day,
    "add_days": formula_add_days, "format_date": formula_format_date, "parse_date": formula_parse_date,
}


def validate_json_schema(value: Any, schema: dict[str, Any]) -> None:
    if not schema:
        return
    try:
        from jsonschema import validate
        from jsonschema.exceptions import ValidationError
        try:
            validate(value, schema)
        except ValidationError as exc:
            raise ValueError(exc.message) from exc
    except ImportError:
        if schema.get("type") == "object" and not isinstance(value, dict):
            raise ValueError("Ожидался object")
        for required in schema.get("required", []):
            if not isinstance(value, dict) or required not in value:
                raise ValueError(f"Отсутствует обязательное поле: {required}")


def parse_json_response(text: str) -> Any:
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"```(?:json)?\s*(.*?)```", text, re.S)
        if match:
            return json.loads(match.group(1))
        start = min((position for position in (text.find("{"), text.find("[")) if position >= 0), default=-1)
        if start >= 0:
            return json.loads(text[start:])
        raise


def llm_output(parsed: Any, model: str, usage: dict[str, Any]) -> dict[str, Any]:
    records = parsed.get("records") if isinstance(parsed, dict) and "records" in parsed else parsed
    if not isinstance(records, list):
        records = [records]
    return {"parsed_response": parsed, "records": records, "model": model, "usage": usage}


def parse_page_selection(specification: str, page_count: int) -> list[int]:
    if not specification:
        return list(range(page_count))
    result: set[int] = set()
    for part in specification.split(","):
        part = part.strip()
        if "-" in part:
            start, end = (int(value) for value in part.split("-", 1))
            result.update(range(max(start - 1, 0), min(end, page_count)))
        elif part.isdigit() and 1 <= int(part) <= page_count:
            result.add(int(part) - 1)
    return sorted(result)


def dedupe_headers(headers: list[str]) -> list[str]:
    output: list[str] = []
    counts: dict[str, int] = {}
    for index, header in enumerate(headers):
        base = header.strip() or f"column_{index + 1}"
        counts[base] = counts.get(base, 0) + 1
        output.append(base if counts[base] == 1 else f"{base}_{counts[base]}")
    return output


def stringify(value: Any) -> str:
    return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False, default=str)


def stringify_cell(value: Any) -> Any:
    return json.dumps(value, ensure_ascii=False, default=str) if isinstance(value, (dict, list)) else value
