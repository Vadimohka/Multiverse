import csv
import io
import json
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from openpyxl import load_workbook
from sqlalchemy.orm import Session

from app.audit import audit
from app.database import get_db
from app.dependencies import require_roles
from app.models import RawDocument, Source, User
from app.services.artifact_storage import ArtifactStorage
from app.services.authorization import require_project

router = APIRouter(prefix="/documents", tags=["Документы"])
MAX_FILE_SIZE = 100 * 1024 * 1024


async def read_limited(file: UploadFile) -> bytes:
    data = await file.read(MAX_FILE_SIZE + 1)
    if len(data) > MAX_FILE_SIZE: raise HTTPException(status_code=413, detail="Файл превышает 100 МБ")
    return data


@router.post("/parse")
async def parse_document(file: UploadFile = File(...), _: User = Depends(require_roles("ADMINISTRATOR", "DEVELOPER", "OPERATOR"))) -> dict[str, Any]:
    data = await read_limited(file); suffix = Path(file.filename or "").suffix.lower()
    try:
        if suffix == ".csv":
            text = data.decode("utf-8-sig"); dialect = csv.Sniffer().sniff(text[:4096]); rows = list(csv.DictReader(io.StringIO(text), dialect=dialect)); return {"type": "CSV", "records": rows, "count": len(rows), "evidence": {"filename": file.filename}}
        if suffix == ".xlsx":
            wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True); output = {}
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True)); headers = [str(x) if x is not None else f"column_{i+1}" for i,x in enumerate(rows[0])] if rows else []
                output[ws.title] = [dict(zip(headers,row,strict=False)) for row in rows[1:] if any(x is not None for x in row)]
            return {"type": "XLSX", "sheets": output, "evidence": {"filename": file.filename}}
        if suffix == ".docx":
            from docx import Document
            doc = Document(io.BytesIO(data)); tables=[]
            for table in doc.tables: tables.append([[cell.text for cell in row.cells] for row in table.rows])
            return {"type": "DOCX", "paragraphs": [p.text for p in doc.paragraphs], "tables": tables, "evidence": {"filename": file.filename}}
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader=PdfReader(io.BytesIO(data)); pages=[{"page": i+1, "text": page.extract_text() or ""} for i,page in enumerate(reader.pages)]
            return {"type": "PDF", "pages": pages, "page_count": len(pages), "evidence": {"filename": file.filename}}
        if suffix == ".json": return {"type": "JSON", "data": json.loads(data), "evidence": {"filename": file.filename}}
        raise HTTPException(status_code=415, detail="Поддерживаются PDF, DOCX, XLSX, CSV и JSON")
    except HTTPException: raise
    except Exception as exc: raise HTTPException(status_code=422, detail=f"Ошибка разбора документа: {exc}") from exc


@router.post("/upload-source", status_code=201)
async def upload_source_document(
    project_id: str = Form(...),
    name: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("ADMINISTRATOR", "DEVELOPER")),
) -> dict[str, Any]:
    """Store an uploaded document as a reusable source and raw artifact."""
    require_project(db, user, project_id)
    data = await read_limited(file)
    filename = Path(file.filename or "document").name
    if Path(filename).suffix.lower() not in {".pdf", ".docx", ".xlsx", ".csv", ".json"}:
        raise HTTPException(status_code=415, detail="Поддерживаются PDF, DOCX, XLSX, CSV и JSON")
    source = Source(project_id=project_id, name=name, source_type="DOCUMENT", entry_url=f"document://{filename}", fetch_mode="DOCUMENT", settings={})
    db.add(source)
    db.flush()
    stored = await ArtifactStorage().put_bytes("raw", f"sources/{source.id}/{filename}", data, file.content_type or "application/octet-stream", {"source_id": source.id})
    source.settings = {
        "document_storage_key": stored["storage_key"],
        "document_storage_backend": stored.get("storage_backend", "S3"),
        "document_bucket": stored.get("bucket", "raw"),
        "document_filename": filename,
        "document_content_type": file.content_type or "application/octet-stream",
    }
    db.add(RawDocument(source_id=source.id, url=source.entry_url, content_type=source.settings["document_content_type"], sha256=stored["sha256"], storage_key=stored["storage_key"], metadata_json={"filename": filename, **stored}))
    audit(db, user.id, "CREATE", "source", source.id, after={"name": source.name, "source_type": source.source_type, "filename": filename})
    db.commit()
    db.refresh(source)
    return {"id": source.id, "project_id": source.project_id, "name": source.name, "source_type": source.source_type, "entry_url": source.entry_url, "fetch_mode": source.fetch_mode, "settings": source.settings}
